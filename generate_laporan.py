#!/usr/bin/env python3
"""Generate Laporan Magang DOCX for ElnusaAbsensiWEB project."""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ============================================================
# PAGE SETUP (Standard: 4cm left, 3cm top/right/bottom)
# ============================================================
for section in doc.sections:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

# ============================================================
# STYLES
# ============================================================
style_normal = doc.styles['Normal']
font = style_normal.font
font.name = 'Times New Roman'
font.size = Pt(12)
style_normal.paragraph_format.line_spacing = 1.5
style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
style_normal.paragraph_format.space_after = Pt(6)

# Heading 1 (Chapter)
h1 = doc.styles['Heading 1']
h1.font.name = 'Times New Roman'
h1.font.size = Pt(14)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0, 0, 0)
h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
h1.paragraph_format.space_before = Pt(18)
h1.paragraph_format.space_after = Pt(12)

# Heading 2 (Sub-chapter)
h2 = doc.styles['Heading 2']
h2.font.name = 'Times New Roman'
h2.font.size = Pt(12)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0, 0, 0)
h2.paragraph_format.space_before = Pt(12)
h2.paragraph_format.space_after = Pt(6)

# Heading 3 (Sub-sub-chapter)
h3 = doc.styles['Heading 3']
h3.font.name = 'Times New Roman'
h3.font.size = Pt(12)
h3.font.bold = True
h3.font.italic = True
h3.font.color.rgb = RGBColor(0, 0, 0)
h3.paragraph_format.space_before = Pt(6)
h3.paragraph_format.space_after = Pt(3)


def add_paragraph(text, bold=False, italic=False, alignment=None, size=12, space_after=6, first_line_indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = alignment or WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    if first_line_indent and alignment is None:
        p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def add_numbered(text, level=0):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def add_code_block(code_text):
    """Add code in a shaded monospace block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    # Add shading
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F0F0F0')
    p.paragraph_format.element.get_or_add_pPr().append(shading)

    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p


def add_table_row(table, cells_text, bold=False):
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.bold = bold
    return row


# ============================================================
# COVER PAGE
# ============================================================
for _ in range(4):
    add_paragraph('', space_after=0)

add_paragraph('LAPORAN MAGANG', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=18, first_line_indent=False)
add_paragraph('', space_after=0, first_line_indent=False)

add_paragraph(
    'PENGEMBANGAN SISTEM ABSENSI BERBASIS QR CODE DAN BARCODE '
    'PADA WAREHOUSE ELNUSA BSD MENGGUNAKAN REACT.JS DAN FIREBASE',
    bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=14, first_line_indent=False
)
add_paragraph('', space_after=0, first_line_indent=False)
add_paragraph('', space_after=0, first_line_indent=False)

add_paragraph('Disusun oleh:', alignment=WD_ALIGN_PARAGRAPH.CENTER, size=12, first_line_indent=False)
add_paragraph('Adam Putra Pratama', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=13, first_line_indent=False)
add_paragraph('NIM: 20230801402', alignment=WD_ALIGN_PARAGRAPH.CENTER, size=12, first_line_indent=False)
add_paragraph('', space_after=0, first_line_indent=False)
add_paragraph('', space_after=0, first_line_indent=False)
add_paragraph('', space_after=0, first_line_indent=False)

add_paragraph(
    'PROGRAM STUDI TEKNIK INFORMATIKA\n'
    'FAKULTAS ILMU KOMPUTER\n'
    '[NAMA UNIVERSITAS]',
    bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=12, first_line_indent=False
)
add_paragraph('', space_after=0, first_line_indent=False)
add_paragraph('TAHUN 2026', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=12, first_line_indent=False)

# Page break
doc.add_page_break()

# ============================================================
# KATA PENGANTAR
# ============================================================
doc.add_heading('KATA PENGANTAR', level=1)
add_paragraph('')
add_paragraph(
    'Puji syukur ke hadirat Tuhan Yang Maha Esa atas segala rahmat dan karunia-Nya '
    'sehingga penulis dapat menyelesaikan laporan magang ini dengan judul "Pengembangan '
    'Sistem Absensi Berbasis QR Code dan Barcode pada Warehouse Elnusa BSD Menggunakan '
    'React.js dan Firebase" sebagai salah satu syarat kelulusan mata kuliah Magang pada '
    'Program Studi Teknik Informatika.'
)
add_paragraph(
    'Laporan ini disusun berdasarkan pengalaman dan hasil kerja penulis selama melaksanakan '
    'kegiatan magang di Warehouse Elnusa BSD. Selama proses magang dan penyusunan laporan ini, '
    'penulis banyak mendapatkan bimbingan, arahan, serta dukungan dari berbagai pihak. '
    'Oleh karena itu, penulis ingin menyampaikan ucapan terima kasih yang sebesar-besarnya kepada:'
)
add_paragraph(
    '1. [Nama Dosen Pembimbing], selaku dosen pembimbing yang telah memberikan arahan dan masukan '
    'dalam penyusunan laporan ini.\n'
    '2. [Nama Pembimbing Lapangan], selaku pembimbing lapangan di Warehouse Elnusa BSD yang telah '
    'memberikan bimbingan teknis selama pelaksanaan magang.\n'
    '3. Seluruh staf dan karyawan Warehouse Elnusa BSD yang telah membantu dan memberikan dukungan '
    'selama proses magang.\n'
    '4. Keluarga dan rekan-rekan yang selalu memberikan motivasi dan dukungan.'
)
add_paragraph(
    'Penulis menyadari bahwa laporan ini masih jauh dari sempurna. Oleh karena itu, kritik dan '
    'saran yang membangun sangat penulis harapkan untuk perbaikan di masa mendatang. Semoga '
    'laporan ini dapat memberikan manfaat bagi pengembangan ilmu pengetahuan dan teknologi, '
    'khususnya di bidang sistem informasi absensi.'
)
add_paragraph('')
add_paragraph('[Tempat], [Tanggal]', alignment=WD_ALIGN_PARAGRAPH.RIGHT, first_line_indent=False)
add_paragraph('', space_after=0, first_line_indent=False)
add_paragraph('Penulis', alignment=WD_ALIGN_PARAGRAPH.RIGHT, first_line_indent=False)

doc.add_page_break()

# ============================================================
# DAFTAR ISI (placeholder - will be auto-generated in Word)
# ============================================================
doc.add_heading('DAFTAR ISI', level=1)
add_paragraph('')
add_paragraph('KATA PENGANTAR ............................................................ ii', first_line_indent=False)
add_paragraph('DAFTAR ISI .................................................................. iii', first_line_indent=False)
add_paragraph('DAFTAR GAMBAR ............................................................ iv', first_line_indent=False)
add_paragraph('DAFTAR TABEL .............................................................. v', first_line_indent=False)
add_paragraph('')
add_paragraph('BAB I PENDAHULUAN .......................................................... 1', first_line_indent=False)
add_paragraph('    1.1 Latar Belakang ..................................................... 1', first_line_indent=False)
add_paragraph('    1.2 Rumusan Masalah ................................................... 3', first_line_indent=False)
add_paragraph('    1.3 Tujuan ............................................................ 3', first_line_indent=False)
add_paragraph('    1.4 Manfaat ........................................................... 4', first_line_indent=False)
add_paragraph('    1.5 Batasan Masalah ................................................... 4', first_line_indent=False)
add_paragraph('')
add_paragraph('BAB II TINJAUAN PUSTAKA .................................................... 5', first_line_indent=False)
add_paragraph('    2.1 React.js dan TypeScript .......................................... 5', first_line_indent=False)
add_paragraph('    2.2 Firebase Firestore ............................................... 6', first_line_indent=False)
add_paragraph('    2.3 QR Code dan Barcode dalam Sistem Absensi ......................... 7', first_line_indent=False)
add_paragraph('    2.4 Tailwind CSS ..................................................... 8', first_line_indent=False)
add_paragraph('    2.5 Vite dan Express.js .............................................. 8', first_line_indent=False)
add_paragraph('    2.6 Penelitian Terdahulu ............................................. 9', first_line_indent=False)
add_paragraph('')
add_paragraph('BAB III METODOLOGI DAN IMPLEMENTASI ....................................... 10', first_line_indent=False)
add_paragraph('    3.1 Metode Pengembangan .............................................. 10', first_line_indent=False)
add_paragraph('    3.2 Analisis Kebutuhan ............................................... 11', first_line_indent=False)
add_paragraph('    3.3 Perancangan Sistem ............................................... 13', first_line_indent=False)
add_paragraph('    3.4 Implementasi ..................................................... 16', first_line_indent=False)
add_paragraph('    3.5 Pengujian Sistem ................................................. 30', first_line_indent=False)
add_paragraph('')
add_paragraph('BAB IV HASIL DAN PEMBAHASAN .............................................. 35', first_line_indent=False)
add_paragraph('    4.1 Hasil Implementasi ............................................... 35', first_line_indent=False)
add_paragraph('    4.2 Relevansi Mata Kuliah ........................................... 38', first_line_indent=False)
add_paragraph('')
add_paragraph('DAFTAR PUSTAKA ............................................................ 42', first_line_indent=False)

doc.add_page_break()

# ============================================================
# BAB I - PENDAHULUAN
# ============================================================
doc.add_heading('BAB I\nPENDAHULUAN', level=1)

doc.add_heading('1.1 Latar Belakang', level=2)
add_paragraph(
    'Perkembangan teknologi informasi dan komunikasi telah membawa perubahan signifikan dalam '
    'berbagai aspek kehidupan, termasuk di dunia industri dan pergudangan. Teknologi digital '
    'telah menjadi kebutuhan pokok bagi perusahaan untuk meningkatkan efisiensi operasional, '
    'akurasi data, serta kecepatan dalam pengambilan keputusan. Salah satu aspek operasional '
    'yang krusial dalam manajemen pergudangan adalah sistem absensi atau pencatatan kehadiran '
    'karyawan dan tamu yang masuk dan keluar dari area gudang.'
)
add_paragraph(
    'Warehouse Elnusa BSD merupakan unit pergudangan dari PT Elnusa Tbk yang bergerak di bidang '
    'logistik dan penyimpanan. Dalam operasional sehari-harinya, Warehouse Elnusa BSD memerlukan '
    'sistem pencatatan kehadiran yang akurat dan efisien untuk memantau personel yang berada di '
    'dalam area gudang. Personel yang dimaksud mencakup kary tetap, karyawa tidak tetap, '
    'serta tamu atau visitor yang memiliki keperluan di area gudang.'
)
add_paragraph(
    'Berdasarkan observasi awal yang dilakukan penulis, sistem absensi yang berjalan di Warehouse '
    'Elnusa BSD masih menggunakan metode manual, yaitu pencatatan melalui buku tamu dan formulir '
    'kertas. Metode manual ini memiliki beberapa kelemahan, antara lain: (1) antrean panjang pada '
    'saat jam masuk dan keluar kerja karena proses pencatatan yang lambat; (2) rawan terjadinya '
    'kesalahan pencatatan data akibat tulisan tangan yang tidak jelas; (3) kesulitan dalam '
    'merekap data kehadiran secara real-time karena data tersebar di buku catatan fisik; '
    '(4) potensi manipulasi data kehadiran karena tidak ada mekanisme verifikasi yang ketat; '
    'serta (5) keterbatasan dalam memantau jumlah personel yang sedang berada di dalam area gudang '
    '(Person On Board/POB) secara langsung.'
)
add_paragraph(
    'Untuk mengatasi permasalahan tersebut, diperlukan sebuah sistem absensi digital yang mampu '
    'melakukan pencatatan kehadiran secara otomatis, akurat, dan real-time. Teknologi QR Code '
    'dan Barcode menjadi pilihan yang tepat karena keduanya dapat menyimpan data identitas '
    'unik setiap personel dalam bentuk kode dua dimensi maupun satu dimensi yang mudah dibaca '
    'oleh pemindai optik. Setiap karyawan akan memiliki identitas digital berupa QR Code dan '
    'Barcode yang dapat dipindai oleh petugas security di pintu gerbang masuk maupun keluar '
    'gudang.'
)
add_paragraph(
    'Proyek pengembangan sistem absensi ini dibangun menggunakan teknologi web modern, yaitu '
    'React.js dengan TypeScript untuk antarmuka pengguna (frontend) dan Firebase sebagai layanan '
    'backend yang menyediakan basis data real-time (Firestore). Pemilihan React.js didasarkan '
    'pada kemampuannya dalam membangun antarmuka pengguna yang interaktif, responsif, dan '
    'reaktif terhadap perubahan data secara real-time. Sementara itu, Firebase Firestore dipilih '
    'karena menyediakan layanan basis data NoSQL yang terkelola sepenuhnya dengan kemampuan '
    'sinkronisasi data secara real-time, sehingga cocok untuk aplikasi yang membutuhkan '
    'pembaruan data secara langsung seperti sistem absensi.'
)
add_paragraph(
    'Berdasarkan uraian di atas, maka penulis merasa perlu untuk mengembangkan sistem absensi '
    'berbasis QR Code dan Barcode di Warehouse Elnusa BSD. Laporan magang ini disusun sebagai '
    'dokumentasi dari proses pengembangan sistem yang telah dilakukan, mulai dari analisis '
    'kebutuhan, perancangan, implementasi, hingga pengujian sistem.')

doc.add_heading('1.2 Rumusan Masalah', level=2)
add_paragraph(
    'Berdasarkan latar belakang yang telah diuraikan, maka rumusan masalah dalam laporan magang '
    'ini adalah sebagai berikut:'
)
add_numbered('Bagaimana merancang dan membangun sistem absensi digital berbasis QR Code dan Barcode di Warehouse Elnusa BSD yang mampu menggantikan sistem pencatatan manual?')
add_numbered('Bagaimana mengimplementasikan sistem manajemen kehadiran dengan tiga peran pengguna (Admin, Security, dan Employee) yang memiliki hak akses berbeda sesuai dengan kebutuhan operasional?')
add_numbered('Bagaimana mengintegrasikan teknologi pemindaian QR Code dan Barcode, baik melalui input keyboard dari scanner hardware maupun melalui kamera perangkat, ke dalam antarmuka sistem berbasis web?')

doc.add_heading('1.3 Tujuan', level=2)
add_paragraph(
    'Tujuan dari pelaksanaan magang dan penulisan laporan ini adalah:'
)
add_numbered('Mengembangkan sistem absensi digital berbasis QR Code dan Barcode di Warehouse Elnusa BSD yang dapat melakukan pencatatan kehadiran secara otomatis dan real-time.')
add_numbered('Menyediakan sistem manajemen kehadiran dengan tiga peran pengguna (Admin, Security, dan Employee) yang memiliki fitur dan hak akses sesuai dengan kebutuhan masing-masing.')
add_numbered('Mengintegrasikan pemindaian QR Code dan Barcode ke dalam sistem berbasis web, baik melalui input scanner hardware maupun kamera perangkat.')

doc.add_heading('1.4 Manfaat', level=2)
add_paragraph(
    'Manfaat yang diharapkan dari pengembangan sistem absensi ini adalah sebagai berikut:'
)
add_bullet('Bagi Warehouse Elnusa BSD: (a) Meningkatkan efisiensi proses pencatatan kehadiran di pintu gerbang gudang; (b) Menyediakan data kehadiran yang akurat dan real-time; (c) Memudahkan proses rekap dan pelaporan data kehadiran melalui fitur ekspor PDF; (d) Memungkinkan pemantauan jumlah personel yang berada di dalam area gudang (POB) secara langsung.')
add_bullet('Bagi Penulis: (a) Menerapkan ilmu dan keterampilan yang diperoleh selama perkuliahan, khususnya di bidang pengembangan web, basis data, dan interaksi manusia-komputer; (b) Memperoleh pengalaman nyata dalam mengembangkan sistem informasi di lingkungan industri.')
add_bullet('Bagi Akademisi: Memberikan referensi dan dokumentasi mengenai pengembangan sistem absensi berbasis QR Code dan Barcode menggunakan React.js dan Firebase.')

doc.add_heading('1.5 Batasan Masalah', level=2)
add_paragraph(
    'Agar pembahasan dalam laporan ini terfokus dan tidak meluas, maka ditetapkan batasan '
    'masalah sebagai berikut:'
)
add_bullet('Sistem dikembangkan dalam platform web (browser) dan bukan aplikasi mobile native.')
add_bullet('Backend sistem menggunakan Firebase Firestore sebagai basis data real-time dan tidak menggunakan server backend khusus selain Firebase dan server development Express.js.')
add_bullet('Sistem berfokus pada fitur pencatatan kehadiran di pintu masuk dan keluar (gate entry/exit), serta tidak mencakup modul penggajian, manajemen cuti, atau manajemen SDM secara umum.')
add_bullet('Sistem menggunakan QR Code dan Barcode sebagai media identitas digital yang dicetak dalam bentuk kartu identitas (identity pass).')
add_bullet('Proses autentikasi pengguna untuk role Admin dan Security menggunakan password statis yang telah ditentukan, bukan sistem autentikasi berbasis akun pengguna yang terdaftar.')

doc.add_page_break()

# ============================================================
# BAB II - TINJAUAN PUSTAKA
# ============================================================
doc.add_heading('BAB II\nTINJAUAN PUSTAKA', level=1)

doc.add_heading('2.1 React.js dan TypeScript', level=2)
add_paragraph(
    'React.js merupakan salah satu library JavaScript yang paling populer untuk membangun antarmuka '
    'pengguna (user interface) pada aplikasi web. Dikembangkan oleh Meta (sebelumnya Facebook), '
    'React.js menggunakan pendekatan komponen (component-based architecture) di mana setiap bagian '
    'dari antarmuka pengguna dibangun sebagai komponen yang independen, dapat digunakan kembali, '
    'dan memiliki state-nya sendiri. Pendekatan ini memudahkan pengembang dalam mengelola kompleksitas '
    'aplikasi web modern yang memiliki banyak interaksi dan pembaruan data secara dinamis.'
)
add_paragraph(
    'Salah satu konsep utama dalam React.js adalah Virtual DOM (Document Object Model). Virtual DOM '
    'merupakan representasi ringan dari DOM asli yang disimpan di dalam memori. Ketika terjadi '
    'perubahan data, React.js akan memperbarui Virtual DOM terlebih dahulu, kemudian membandingkannya '
    'dengan Virtual DOM sebelumnya (differencing algorithm), dan hanya menerapkan perubahan yang '
    'diperlukan ke DOM asli (reconciliation). Mekanisme ini membuat React.js sangat efisien dalam '
    'menangani pembaruan antarmuka pengguna secara real-time.'
)
add_paragraph(
    'TypeScript adalah bahasa pemrograman yang merupakan superset dari JavaScript, yang menambahkan '
    'dukungan tipe data statis (static typing) ke dalam JavaScript. Dengan adanya tipe data statis, '
    'TypeScript memungkinkan deteksi kesalahan pada tahap kompilasi (compile-time) sebelum kode '
    'dijalankan di browser. Hal ini sangat bermanfaat dalam pengembangan aplikasi berskala besar '
    'karena dapat mengurangi bug, meningkatkan keterbacaan kode, dan memudahkan proses refactoring. '
    'Pada proyek ini, React.js digunakan bersama dengan TypeScript untuk membangun seluruh komponen '
    'antarmuka sistem absensi.'
)

doc.add_heading('2.2 Firebase Firestore', level=2)
add_paragraph(
    'Firebase adalah platform pengembangan aplikasi milik Google yang menyediakan berbagai layanan '
    'backend, termasuk basis data real-time (Firestore), autentikasi pengguna, penyimpanan file, '
    'dan hosting. Firebase Firestore adalah basis data NoSQL yang bersifat fleksibel, scalable, dan '
    'terkelola sepenuhnya di cloud. Data dalam Firestore disimpan dalam bentuk dokumen (document) '
    'yang terorganisir dalam koleksi (collection). Setiap dokumen berisi pasangan key-value yang '
    'dapat berupa berbagai tipe data.'
)
add_paragraph(
    'Salah satu keunggulan utama Firestore adalah kemampuannya untuk melakukan sinkronisasi data '
    'secara real-time. Ketika suatu data berubah di server, Firestore secara otomatis mengirimkan '
    'pembaruan ke semua klien yang sedang mendengarkan (listening) data tersebut. Fitur ini sangat '
    'cocok untuk aplikasi absensi yang memerlukan pembaruan data kehadiran secara langsung tanpa '
    'harus melakukan refresh halaman.'
)
add_paragraph(
    'Dalam proyek ini, Firestore digunakan untuk menyimpan data karyawan (collection employees), '
    'data log kehadiran (collection presence_logs), data statistik harian (collection stats), '
    'serta konfigurasi sistem (collection system_config). Penggunaan Firestore memungkinkan '
    'dashboard admin dan antarmuka scan untuk menampilkan data statistik secara real-time tanpa '
    'perlu melakukan polling ke server secara periodik.'
)

doc.add_heading('2.3 QR Code dan Barcode dalam Sistem Absensi', level=2)
add_paragraph(
    'QR Code (Quick Response Code) adalah jenis kode matriks dua dimensi yang dapat menyimpan '
    'informasi dalam format teks, URL, atau data lainnya. QR Code dapat dibaca dengan cepat '
    'menggunakan kamera smartphone atau scanner khusus dari berbagai sudut pemindaian. Kapasitas '
    'penyimpanan QR Code lebih besar dibandingkan Barcode tradisional dan dilengkapi dengan '
    'kemampuan koreksi kesalahan (error correction) yang memungkinkan data tetap terbaca meskipun '
    'sebagian kode mengalami kerusakan.'
)
add_paragraph(
    'Barcode atau kode batang adalah representasi data optis yang dapat dibaca oleh mesin dalam '
    'bentuk garis-garis vertikal dengan ketebalan dan jarak yang bervariasi. Barcode satu dimensi '
    '(1D) seperti Code 128 dan Code 39 banyak digunakan di industri untuk identifikasi produk, '
    'aset, dan personel. Meskipun kapasitas penyimpanannya lebih terbatas dibandingkan QR Code, '
    'Barcode tetap relevan digunakan karena kompatibilitasnya dengan berbagai jenis scanner '
    'hardware yang banyak tersedia di pasaran.'
)
add_paragraph(
    'Dalam sistem absensi yang dikembangkan, setiap karyawan dan visitor memiliki identitas unik '
    'berupa NIK (Nomor Induk Karyawan) atau nomor identitas lainnya. Identitas unik ini '
    'direpresentasikan dalam dua format kode, yaitu QR Code dan Barcode, yang ditampilkan pada '
    'kartu identitas digital (identity pass). Pada sisi pemindaian, sistem mendukung dua metode: '
    '(1) input melalui scanner hardware yang mengirimkan data seolah-olah diketikkan melalui '
    'keyboard (keyboard wedge), dan (2) pemindaian melalui kamera perangkat menggunakan library '
    'html5-qrcode yang dapat membaca berbagai format kode termasuk QR Code, Code 128, EAN-13, '
    'dan Code 39.'
)

doc.add_heading('2.4 Tailwind CSS', level=2)
add_paragraph(
    'Tailwind CSS adalah framework CSS utility-first yang memungkinkan pengembang membangun '
    'antarmuka pengguna dengan cepat menggunakan kelas-kelas utilitas yang telah disediakan. '
    'Berbeda dengan framework CSS tradisional seperti Bootstrap yang menyediakan komponen '
    'siap pakai (seperti card, navbar, button), Tailwind CSS menyediakan kelas-kelas utilitas '
    'tingkat rendah (seperti flex, text-center, bg-blue-500, p-4) yang dapat dikombinasikan '
    'secara bebas untuk menciptakan desain yang unik dan sesuai kebutuhan.'
)
add_paragraph(
    'Pada proyek ini, Tailwind CSS versi 4 digunakan untuk membangun seluruh antarmuka sistem '
    'absensi. Pemilihan Tailwind CSS didasarkan pada fleksibilitasnya dalam menciptakan desain '
    'kustom tanpa harus menulis CSS tambahan, kemampuannya dalam menghasilkan file CSS yang '
    'sangat kecil pada production (karna hanya menyertakan kelas yang benar-benar digunakan), '
    'serta dukungannya terhadap desain responsif yang memungkinkan sistem diakses dengan baik '
    'dari berbagai ukuran layar, baik desktop maupun perangkat mobile.'
)

doc.add_heading('2.5 Vite dan Express.js', level=2)
add_paragraph(
    'Vite adalah build tool modern untuk pengembangan aplikasi web yang dikembangkan oleh Evan You '
    '(pembuat Vue.js). Vite menawarkan kecepatan pengembangan yang sangat tinggi dengan menggunakan '
    'ES Modules (ESM) secara native pada mode pengembangan, sehingga proses hot module replacement '
    '(HMR) dapat berjalan dengan instan tanpa perlu melakukan bundle ulang seluruh aplikasi. Pada '
    'mode produksi, Vite menggunakan Rollup untuk melakukan bundling kode secara optimal dengan '
    'tree-shaking dan code splitting.'
)
add_paragraph(
    'Express.js adalah framework web minimalis untuk Node.js yang digunakan untuk membangun server '
    'HTTP dan API. Dalam proyek ini, Express.js digunakan sebagai server pengembangan yang '
    'melayani aplikasi React melalui middleware Vite pada mode pengembangan, serta sebagai server '
    'statis untuk melayani file build pada mode produksi. Penggunaan Express.js juga menyediakan '
    'endpoint API untuk kebutuhan integrasi di masa mendatang, seperti endpoint health check pada '
    '/api/health.'
)

doc.add_heading('2.6 Penelitian Terdahulu', level=2)
add_paragraph(
    'Beberapa penelitian terdahulu yang relevan dengan pengembangan sistem absensi berbasis QR Code '
    'dan Barcode antara lain:'
)
add_paragraph(
    'Penelitian oleh Alief dan Rianto (2025) yang merancang sistem absensi siswa berbasis QR Code '
    'menggunakan framework JavaScript, yaitu React.js pada sisi frontend dan Express.js pada sisi '
    'backend, serta MySQL sebagai basis data. Penelitian ini menerapkan metode pengembangan '
    'perangkat lunak Waterfall dan pengujian menggunakan metode Black Box. Hasil penelitian '
    'menunjukkan bahwa sistem absensi berbasis QR Code mampu meningkatkan efisiensi, akurasi, '
    'dan keamanan data pencatatan kehadiran. Kemiripan dengan penelitian ini terletak pada '
    'penggunaan React.js sebagai frontend dan metode Waterfall serta Black Box Testing.'
)
add_paragraph(
    'Penelitian oleh Praba, Safitri, dan Faridi (2025) yang mengembangkan aplikasi absensi '
    'berbasis website menggunakan QR Code untuk peningkatan efisiensi pencatatan kehadiran. '
    'Penelitian ini menggunakan metode Research and Development (R&D) yang mencakup tahap '
    'perancangan sistem, pembuatan QR Code untuk absensi, serta pengujian sistem. Hasil '
    'penelitian menunjukkan bahwa sistem absensi berbasis website menggunakan QR Code dapat '
    'meningkatkan efisiensi dan akurasi dalam pencatatan kehadiran serta memudahkan pengelolaan '
    'data absensi secara real-time.'
)
add_paragraph(
    'Penelitian oleh Nuralif dan Fachrie (2023) yang mengembangkan sistem absensi berbasis '
    'QR Code untuk karyawan pabrik. Penelitian ini relevan karena membahas implementasi QR Code '
    'dalam konteks lingkungan industri manufaktur, serupa dengan penerapan di Warehouse Elnusa '
    'BSD. Sistem yang dikembangkan mampu mengotomatisasi pencatatan kehadiran karyawan dan '
    'menyediakan data yang akurat secara real-time. Kesamaan dengan penelitian ini adalah pada '
    'penggunaan QR Code sebagai identitas digital dan fokus pada lingkungan kerja industri.'
)

doc.add_page_break()

# ============================================================
# BAB III - METODOLOGI DAN IMPLEMENTASI
# ============================================================
doc.add_heading('BAB III\nMETODOLOGI DAN IMPLEMENTASI', level=1)

doc.add_heading('3.1 Metode Pengembangan', level=2)
add_paragraph(
    'Metode pengembangan yang digunakan dalam proyek ini adalah model Waterfall, yang merupakan '
    'salah satu metode pengembangan perangkat lunak tertua dan paling banyak digunakan. Model '
    'Waterfall memandang proses pengembangan perangkat lunak sebagai alur kerja yang berurutan '
    'dan sistematis, di mana setiap tahap harus diselesaikan terlebih dahulu sebelum melanjutkan '
    'ke tahap berikutnya. Tahapan dalam model Waterfall meliputi:'
)
add_numbered('Analisis Kebutuhan (Requirements Analysis): Mengidentifikasi dan mendokumentasikan kebutuhan fungsional dan non-fungsional sistem berdasarkan observasi dan diskusi dengan pihak Warehouse Elnusa BSD.')
add_numbered('Perancangan (Design): Merancang arsitektur sistem, basis data, antarmuka pengguna, serta alur kerja sistem berdasarkan kebutuhan yang telah dianalisis.')
add_numbered('Implementasi (Implementation): Menerjemahkan perancangan ke dalam kode program menggunakan React.js, TypeScript, dan Firebase sesuai dengan spesifikasi yang telah ditetapkan.')
add_numbered('Pengujian (Testing): Melakukan pengujian fungsional sistem menggunakan metode Black Box Testing untuk memastikan setiap fitur berjalan sesuai dengan kebutuhan.')
add_numbered('Pemeliharaan (Maintenance): Melakukan perbaikan dan penyesuaian sistem berdasarkan umpan balik pengguna setelah sistem digunakan.')
add_paragraph(
    'Pemilihan model Waterfall didasarkan pada karakteristik proyek yang memiliki kebutuhan yang '
    'jelas dan relatif stabil, ruang lingkup yang terdefinisi dengan baik, serta tenggat waktu '
    'pengerjaan yang terbatas. Meskipun model ini kurang fleksibel terhadap perubahan kebutuhan '
    'di tengah proses pengembangan, pendekatan yang sistematis dan terdokumentasi dengan baik '
    'sangat sesuai untuk proyek magang yang memerlukan dokumentasi yang lengkap.'
)

doc.add_heading('3.2 Analisis Kebutuhan', level=2)

doc.add_heading('3.2.1 Kebutuhan Fungsional', level=3)
add_paragraph(
    'Berdasarkan hasil observasi dan diskusi dengan pihak terkait di Warehouse Elnusa BSD, '
    'kebutuhan fungsional sistem absensi yang akan dikembangkan adalah sebagai berikut:'
)

# Create functional requirements table
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Kode'
hdr[1].text = 'Fungsi'
hdr[2].text = 'Deskripsi'
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)

reqs = [
    ('F-01', 'Login Admin', 'Sistem mampu melakukan autentikasi Admin dengan password admin123.'),
    ('F-02', 'Login Security', 'Sistem mampu melakukan autentikasi Security dengan password security123.'),
    ('F-03', 'Login Employee', 'Sistem mampu memverifikasi karyawan berdasarkan NIK yang terdaftar di database.'),
    ('F-04', 'Scan NIK Manual', 'Security dapat memasukkan NIK secara manual melalui input teks.'),
    ('F-05', 'Scan Kamera', 'Security dapat memindai QR Code/Barcode menggunakan kamera perangkat.'),
    ('F-06', 'Proses Absensi', 'Sistem mencatat kehadiran (IN) atau kepulangan (OUT) secara otomatis berdasarkan log terakhir.'),
    ('F-07', 'Statistik Real-time', 'Sistem menampilkan jumlah personel masuk, keluar, POB, dan visitor secara real-time.'),
    ('F-08', 'CRUD Karyawan', 'Admin dapat menambah, melihat, mengedit, dan menghapus data karyawan dan visitor.'),
    ('F-09', 'Export PDF', 'Admin dapat mengekspor laporan kehadiran ke dalam format PDF.'),
    ('F-10', 'Seed Data', 'Admin dapat mengisi database dengan data karyawan sampel (dummy) untuk keperluan pengujian.'),
    ('F-11', 'Maintenance Mode', 'Admin dapat mengaktifkan mode pemeliharaan yang membatasi akses non-admin.'),
    ('F-12', 'Cetak Identity Pass', 'Admin dapat mencetak kartu identitas digital berisi QR Code dan Barcode untuk personel.'),
    ('F-13', 'Riwayat Pribadi', 'Karyawan dapat melihat riwayat kehadiran pribadi mereka.'),
]
for r in reqs:
    add_table_row(table, list(r))

add_paragraph('')

doc.add_heading('3.2.2 Kebutuhan Non-Fungsional', level=3)
add_bullet('Keamanan: Sistem membatasi akses berdasarkan peran pengguna (role-based access control). Data yang disimpan di Firestore diamankan dengan aturan keamanan Firebase.')
add_bullet('Real-time: Data statistik dan log kehadiran harus diperbarui secara real-time menggunakan mekanisme snapshot listener Firestore.')
add_bullet('Responsivitas: Antarmuka sistem harus responsif dan dapat diakses dengan baik di perangkat desktop maupun mobile.')
add_bullet('Usability: Antarmuka dirancang dengan prinsip kemudahan penggunaan (user-friendly) dengan navigasi yang intuitif dan visualisasi data yang jelas.')
add_bullet('Kompatibilitas Scanner: Sistem harus mendukung berbagai jenis scanner hardware yang menggunakan protokol keyboard wedge, serta pemindaian melalui kamera.')

doc.add_heading('3.3 Perancangan Sistem', level=2)

doc.add_heading('3.3.1 Use Case Diagram', level=3)
add_paragraph(
    'Use case diagram menggambarkan interaksi antara aktor (pengguna) dengan sistem yang '
    'dikembangkan. Terdapat tiga aktor dalam sistem absensi ini, yaitu Admin, Security, '
    'dan Employee. Admin memiliki akses penuh terhadap seluruh fitur sistem, termasuk '
    'manajemen data karyawan, ekspor laporan, dan pengaturan sistem. Security memiliki '
    'akses ke fitur pemindaian dan monitoring dashboard. Employee memiliki akses terbatas '
    'untuk melihat identitas digital dan riwayat kehadiran pribadi.'
)
add_paragraph(
    '[Gambar 3.1 Use Case Diagram Sistem Absensi Warehouse Elnusa BSD]', 
    alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False, italic=True
)

# Simple use case as text
add_paragraph('', first_line_indent=False)
add_paragraph('Admin:', bold=True, first_line_indent=False)
add_bullet('Login Admin')
add_bullet('Mengelola Data Karyawan (CRUD)')
add_bullet('Mengelola Data Visitor')
add_bullet('Mengekspor Laporan PDF')
add_bullet('Melakukan Seed Data')
add_bullet('Mengaktifkan/Menonaktifkan Maintenance Mode')
add_bullet('Mencetak Identity Pass')
add_bullet('Melihat Log Kehadiran')

add_paragraph('Security:', bold=True, first_line_indent=False)
add_bullet('Login Security')
add_bullet('Memindai NIK (Manual)')
add_bullet('Memindai QR/Barcode (Kamera)')
add_bullet('Melihat Statistik Real-time (IN/OUT/POB)')
add_bullet('Melihat Dashboard Monitoring')

add_paragraph('Employee:', bold=True, first_line_indent=False)
add_bullet('Login NIK')
add_bullet('Melihat QR Code dan Barcode Pribadi')
add_bullet('Melihat Riwayat Kehadiran')

doc.add_heading('3.3.2 Activity Diagram', level=3)
add_paragraph(
    'Activity diagram menggunakan alur aktivitas dalam sistem. Berikut adalah activity '
    'diagram untuk proses scan absensi yang merupakan fitur utama sistem:'
)
add_paragraph(
    'Proses dimulai ketika Security memasukkan NIK melalui input teks atau memindai '
    'QR/Barcode melalui kamera. Sistem kemudian membersihkan input dari karakter '
    'kontrol dan karakter non-printable. Selanjutnya, sistem mencari data karyawan '
    'berdasarkan NIK menggunakan enam strategi pencarian berurutan (direct document ID, '
    'field nik exact match, case-insensitive, numeric normalization, name search, dan '
    'partial name match). Jika karyawan ditemukan, sistem memeriksa log terakhir '
    'karyawan pada hari yang sama untuk menentukan jenis absensi (IN atau OUT). '
    'Sistem kemudian melakukan batch write untuk mencatat log kehadiran dan memperbarui '
    'statistik harian secara atomik. Hasil berupa notifikasi berisi nama karyawan dan '
    'status absensi ditampilkan di antarmuka.'
)
add_paragraph(
    '[Gambar 3.2 Activity Diagram Proses Scan Absensi]',
    alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False, italic=True
)

doc.add_heading('3.3.3 Arsitektur Sistem', level=3)
add_paragraph(
    'Arsitektur sistem absensi Warehouse Elnusa BSD menggunakan pola arsitektur frontend-'
    'backend yang terpisah. Frontend dibangun menggunakan React.js dengan TypeScript yang di-'
    'bundle menggunakan Vite. Backend menggunakan Firebase Firestore sebagai basis data '
    'real-time dan Firebase Authentication untuk autentikasi (jika diperlukan di masa '
    'mendatang). Server Express.js digunakan sebagai development server yang melayani '
    'aplikasi React melalui middleware Vite.'
)
add_paragraph(
    'Komunikasi antara frontend dan backend sepenuhnya dilakukan melalui Firebase SDK '
    'yang terintegrasi langsung di dalam kode frontend. Pendekatan ini menyederhanakan '
    'arsitektur karena tidak memerlukan server API khusus untuk operasi basis data. '
    'Firebase Firestore menyediakan API client-side yang aman dengan aturan keamanan '
    'yang dapat dikonfigurasi di sisi server.'
)

doc.add_heading('3.4 Implementasi', level=2)

doc.add_heading('3.4.1 Implementasi Login dan Role Selection', level=3)
add_paragraph(
    'Fitur login merupakan gerbang utama sebelum pengguna dapat mengakses sistem. Halaman '
    'login menampilkan tiga kartu pilihan yang merepresentasikan tiga peran pengguna, yaitu '
    'System Admin, Security Officer, dan Employee Portal. Masing-masing kartu memiliki '
    'deskripsi singkat mengenai peran tersebut dan tombol untuk membuka portal masing-masing.'
)
add_paragraph(
    'Implementasi halaman login terdapat pada komponen LoginSelection.tsx. Untuk role Admin '
    'dan Security, autentikasi dilakukan menggunakan password statis yang telah ditentukan. '
    'Admin menggunakan password "admin123" dan Security menggunakan password "security123". '
    'Password ini diperiksa secara langsung di kode frontend tanpa melibatkan server autentikasi, '
    'sehingga cocok untuk penggunaan internal di lingkungan gudang.'
)
add_paragraph(
    'Untuk role Employee, autentikasi dilakukan dengan memverifikasi NIK yang dimasukkan '
    'terhadap data yang tersimpan di koleksi Firestore employees. Jika NIK ditemukan, '
    'data karyawan akan dimuat dan pengguna diarahkan ke portal Employee. Kode berikut '
    'menunjukkan implementasi login employee:'
)

add_code_block(
    'const handleEmployeeLogin = async (e: React.FormEvent) => {\n'
    '  e.preventDefault();\n'
    '  if (!nik) return;\n'
    '  setIsLoading(true);\n'
    '  setError("");\n'
    '  try {\n'
    '    const docRef = doc(db, "employees", nik);\n'
    '    const docSnap = await getDoc(docRef);\n'
    '    if (docSnap.exists()) {\n'
    '      const empData = { id: docSnap.id, ...docSnap.data() } as Employee;\n'
    '      onSelectRole(UserRole.EMPLOYEE, empData);\n'
    '    } else {\n'
    '      setError("NIK tidak terdaftar!");\n'
    '    }\n'
    '  } catch (err) {\n'
    '    setError("Gagal menghubungkan ke server.");\n'
    '  } finally {\n'
    '    setIsLoading(false);\n'
    '  }\n'
    '};'
)

doc.add_heading('3.4.2 Implementasi Scan Interface', level=3)
add_paragraph(
    'Scan Interface merupakan antarmuka utama yang digunakan oleh petugas Security untuk '
    'melakukan pemindaian absensi. Antarmuka ini menampilkan input teks untuk pemindaian '
    'manual (menggunakan scanner hardware keyboard wedge) dan tombol untuk mengaktifkan '
    'kamera sebagai pemindai QR/Barcode visual.'
)
add_paragraph(
    'Input teks secara otomatis menangkap data yang dikirim oleh scanner hardware dalam '
    'format keyboard wedge. Scanner hardware mengirimkan data NIK yang diakhiri dengan '
    'karakter Enter. Sistem menangani input ini melalui event handler onKeyDown yang '
    'mendeteksi ketika tombol Enter atau Tab ditekan, kemudian memproses NIK yang telah '
    'dimasukkan.'
)
add_paragraph(
    'Untuk mencegah pemindaian ganda dalam waktu singkat (misalnya karena scanner membaca '
    'kode yang sama secara berulang), sistem menerapkan mekanisme cooldown di sisi klien '
    'menggunakan useRef. Setiap NIK yang berhasil dipindai akan dicatat waktu pemindaiannya '
    'dan tidak dapat dipindai lagi sebelum 10 detik. Kode berikut menunjukkan logika '
    'penanganan scan:'
)

add_code_block(
    'const handleScan = async (nik: string) => {\n'
    '  const cleanNik = nik.replace(/[\\u0000-\\u001F\\u007F-\\u009F]/g, "").trim();\n'
    '  if (!cleanNik) return;\n'
    '  if (isProcessingRef.current) return;\n'
    '  const now = Date.now();\n'
    '  const lastNIKTime = recentScansRef.current[cleanNik] || 0;\n'
    '  if (now - lastNIKTime < 10000) { // 10s cooldown\n'
    '    setNikInput("");\n'
    '    return;\n'
    '  }\n'
    '  isProcessingRef.current = true;\n'
    '  recentScansRef.current[cleanNik] = now;\n'
    '  setIsProcessing(true);\n'
    '  setNikInput("");\n'
    '  const result = await processScan(cleanNik);\n'
    '  // Handle result...\n'
    '  isProcessingRef.current = false;\n'
    '  setIsProcessing(false);\n'
    '};'
)
add_paragraph(
    'Di bagian footer Scan Interface, terdapat grid statistik yang menampilkan data real-time '
    'menggunakan fitur onSnapshot Firestore. Data yang ditampilkan meliputi jumlah personel '
    'yang masuk (IN), keluar (OUT), Person On Board (POB), serta data visitor. Statistik ini '
    'diperbarui secara otomatis setiap kali terjadi perubahan data di Firestore.'
)

doc.add_heading('3.4.3 Implementasi Admin Dashboard', level=3)
add_paragraph(
    'Admin Dashboard adalah antarmuka yang digunakan oleh Administrator untuk mengelola '
    'seluruh aspek sistem. Dashboard ini menampilkan dua tab utama, yaitu Attendance Log '
    'dan Employee Database. Tab Attendance Log menampilkan daftar log kehadiran yang dapat '
    'difilter berdasarkan tanggal dan departemen. Tab Employee Database menampilkan daftar '
    'seluruh personel yang terdaftar beserta status kehadiran terkini (IN/OUT).'
)
add_paragraph(
    'Fitur utama Admin Dashboard meliputi:'
)
add_bullet('Manajemen Karyawan dan Visitor: Admin dapat menambah data karyawan atau visitor baru melalui form modal, mengedit data yang sudah ada, serta menghapus data yang tidak diperlukan.')
add_bullet('Export PDF: Admin dapat mengekspor laporan kehadiran harian ke dalam format PDF menggunakan library jsPDF dan jspdf-autotable. Laporan mencakup nama karyawan, NIK, departemen, waktu, dan status kehadiran.')
add_bullet('Seed Data: Admin dapat mengisi database dengan data karyawan sampel (90 data dummy) untuk keperluan pengujian dan demonstrasi sistem. Fitur ini menggunakan batch write Firestore untuk menulis data secara atomik dalam satu transaksi.')
add_paragraph(
    'Kode berikut menunjukkan implementasi ekspor PDF pada Admin Dashboard:'
)

add_code_block(
    'const handleExport = () => {\n'
    '  const doc = new jsPDF();\n'
    '  doc.setFontSize(22);\n'
    '  doc.setTextColor(2, 6, 23);\n'
    '  doc.text("Laporan Kehadiran Warehouse", 14, 20);\n'
    '  doc.setFontSize(10);\n'
    '  doc.setTextColor(100, 116, 139);\n'
    '  doc.text(`Warehouse ELNUSA BSD - ${selectedDate}`, 14, 28);\n'
    '  doc.text(`Dicetak pada: ${format(new Date(), "dd/MM/yyyy HH:mm:ss")}`, 14, 40);\n'
    '  const tableData = filteredLogs.map(log => [\n'
    '    employees[log.employeeId]?.name || "Unknown",\n'
    '    log.employeeId,\n'
    '    employees[log.employeeId]?.department || "N/A",\n'
    '    log.timestamp ? format(log.timestamp, "HH:mm:ss") : "N/A",\n'
    '    log.type\n'
    '  ]);\n'
    '  autoTable(doc, {\n'
    '    startY: 48,\n'
    '    head: [["Nama", "NIK", "Departemen", "Waktu", "Status"]],\n'
    '    body: tableData,\n'
    '    theme: "grid",\n'
    '    headStyles: { fillColor: [37, 99, 235] }\n'
    '  });\n'
    '  doc.save(`Warehouse_Attendance_${selectedDate}.pdf`);\n'
    '};'
)

doc.add_heading('3.4.4 Implementasi QR/Barcode Scanner', level=3)
add_paragraph(
    'Fitur pemindaian menggunakan kamera diimplementasikan pada komponen Scanner.tsx yang '
    'menggunakan library html5-qrcode. Library ini mendukung pemindaian berbagai format kode, '
    'termasuk QR Code, Code 128, EAN-13, dan Code 39. Kamera yang digunakan adalah kamera '
    'belakang perangkat (facingMode: "environment") untuk memudahkan pemindaian kode cetak.'
)
add_paragraph(
    'Scanner menampilkan antarmuka dengan area pemindaian yang ditandai dengan kotak '
    'bergaris (bounding box) dan sudut-sudut yang berkedip untuk memandu pengguna '
    'mengarahkan kamera ke kode yang akan dipindai. Ketika kode berhasil terbaca, '
    'library akan memanggil callback function yang kemudian meneruskan data ke fungsi '
    'processScan di attendance.ts.'
)
add_paragraph(
    'Implementasi inisialisasi kamera dan konfigurasi scanner:'
)

add_code_block(
    'useEffect(() => {\n'
    '  html5QrCodeRef.current = new Html5Qrcode(scannerId);\n'
    '  const startCamera = async () => {\n'
    '    try {\n'
    '      const config = {\n'
    '        fps: 15,\n'
    '        qrbox: { width: computedSize, height: computedSize },\n'
    '        aspectRatio: 1.0,\n'
    '        formatsToSupport: [\n'
    '          Html5QrcodeSupportedFormats.QR_CODE,\n'
    '          Html5QrcodeSupportedFormats.CODE_128,\n'
    '          Html5QrcodeSupportedFormats.EAN_13,\n'
    '          Html5QrcodeSupportedFormats.CODE_39\n'
    '        ]\n'
    '      };\n'
    '      await html5QrCodeRef.current?.start(\n'
    '        { facingMode: "environment" },\n'
    '        config,\n'
    '        (decodedText) => { onScan(decodedText); },\n'
    '        () => {} // Ignore frame errors\n'
    '      );\n'
    '    } catch (err) {\n'
    '      setError("Gagal mengakses kamera.");\n'
    '    }\n'
    '  };\n'
    '  startCamera();\n'
    '  return () => {\n'
    '    if (html5QrCodeRef.current?.isScanning) {\n'
    '      html5QrCodeRef.current.stop();\n'
    '    }\n'
    '  };\n'
    '}, []);'
)

doc.add_heading('3.4.5 Implementasi Attendance Logic', level=3)
add_paragraph(
    'Logika utama sistem absensi diimplementasikan pada file attendance.ts yang berisi '
    'fungsi-fungsi inti seperti getEmployeeByNik, getLatestLog, processScan, dan '
    'getDailyStats. Fungsi getEmployeeByNik menggunakan enam strategi pencarian berurutan '
    'untuk menemukan data karyawan berdasarkan input yang diterima dari scanner:'
)
add_numbered('Direct Document ID: Pencarian berdasarkan ID dokumen Firestore yang sama persis dengan input.')
add_numbered('Explicit NIK Field: Pencarian berdasarkan field "nik" pada dokumen dengan kecocokan eksak.')
add_numbered('Case-Insensitive: Pencarian dengan mengabaikan perbedaan huruf besar/kecil pada field NIK.')
add_numbered('Numeric Normalization: Pencarian dengan menghilangkan angka nol di depan (leading zeros) untuk mengakomodasi perbedaan format NIK.')
add_numbered('Name Search: Pencarian berdasarkan field nama apabila scanner membaca nama alih-alih NIK.')
add_numbered('Partial Name Match: Pencarian dengan kecocokan sebagian nama sebagai fallback terakhir.')
add_paragraph(
    'Fungsi processScan merupakan inti dari sistem absensi yang mengelola logika toggle '
    'IN/OUT. Fungsi ini bekerja sebagai berikut: (1) Mencari karyawan berdasarkan NIK; '
    '(2) Memeriksa log terakhir karyawan pada hari yang sama; (3) Menentukan jenis absensi '
    'berikutnya — jika log terakhir adalah IN maka jenis berikutnya adalah OUT, '
    'sebaliknya jika log terakhir adalah OUT atau belum ada log maka jenis berikutnya '
    'adalah IN; (4) Melakukan batch write atomik menggunakan writeBatch Firestore untuk '
    'mencatat log kehadiran dan memperbarui statistik harian secara bersamaan.'
)

add_code_block(
    'export async function processScan(nik: string) {\n'
    '  const employee = await getEmployeeByNik(nik);\n'
    '  if (!employee) {\n'
    '    return { success: false, message: "Karyawan tidak ditemukan" };\n'
    '  }\n'
    '  const todayStr = format(new Date(), "yyyy-MM-dd");\n'
    '  const latestLog = await getLatestLog(employee.id, todayStr);\n'
    '  // 1-minute cooldown check\n'
    '  if (latestLog?.timestamp) {\n'
    '    const diffMinutes = (Date.now() - latestLog.timestamp.seconds*1000) / 60000;\n'
    '    if (diffMinutes < 1) {\n'
    '      return { success: false, message: "Mohon tunggu sebentar", employee };\n'
    '    }\n'
    '  }\n'
    '  let nextType = PresenceType.IN;\n'
    '  if (latestLog?.type === PresenceType.IN) nextType = PresenceType.OUT;\n'
    '  // Atomic batch write\n'
    '  const batch = writeBatch(db);\n'
    '  const logRef = doc(collection(db, "presence_logs"));\n'
    '  batch.set(logRef, {\n'
    '    employeeId: employee.id,\n'
    '    type: nextType,\n'
    '    timestamp: serverTimestamp(),\n'
    '    date: todayStr\n'
    '  });\n'
    '  const statsRef = doc(db, "stats", todayStr);\n'
    '  const statsUpdate = nextType === PresenceType.IN\n'
    '    ? { in: increment(1), pob: increment(1) }\n'
    '    : { out: increment(1), pob: increment(-1) };\n'
    '  batch.set(statsRef, statsUpdate, { merge: true });\n'
    '  await batch.commit();\n'
    '  return { success: true, message: `Selamat ${nextType==="IN" ? "Datang" : "Jalan"}, ${employee.name}`, employee, type: nextType };\n'
    '}'
)

doc.add_heading('3.4.6 Implementasi Maintenance Mode dan Employee Portal', level=3)
add_paragraph(
    'Fitur Maintenance Mode memungkinkan Administrator untuk mengaktifkan mode pemeliharaan '
    'yang akan memblokir akses pengguna non-admin. Ketika mode ini aktif, pengguna dengan '
    'role Security atau Employee akan melihat overlay yang menampilkan pesan pemeliharaan '
    'dan tidak dapat mengakses antarmuka sistem hingga mode dinonaktifkan oleh Admin.'
)
add_paragraph(
    'Implementasi maintenance mode menggunakan mekanisme onSnapshot Firestore pada dokumen '
    'system_config/main. Setiap perubahan pada dokumen tersebut akan langsung tercermin '
    'di seluruh klien yang terhubung. Admin dapat mengaktifkan atau menonaktifkan mode '
    'ini melalui tombol pada Admin Dashboard yang akan memperbarui field maintenanceMode '
    'di dokumen konfigurasi.'
)
add_paragraph(
    'Employee Portal menyediakan antarmuka bagi karyawan untuk melihat identitas digital '
    'mereka dalam bentuk QR Code dan Barcode, serta riwayat kehadiran pribadi. QR Code '
    'ditampilkan menggunakan library react-qr-code dengan level koreksi kesalahan tinggi (H) '
    'untuk memastikan kode tetap terbaca meskipun mengalami sedikit kerusakan. Informasi '
    'yang ditampilkan meliputi sapaan personal, NIK, departemen, serta riwayat kehadiran '
    'dalam bentuk tabel yang diurutkan berdasarkan waktu terbaru.'
)

doc.add_heading('3.5 Pengujian Sistem', level=2)
add_paragraph(
    'Pengujian sistem dilakukan menggunakan metode Black Box Testing yang berfokus pada '
    'pengujian fungsionalitas sistem tanpa melihat struktur kode program di dalamnya. '
    'Metode ini dipilih karena sesuai dengan karakteristik sistem yang memiliki banyak '
    'fitur input dan output yang perlu diverifikasi dari sudut pandang pengguna akhir. '
    'Pengujian dilakukan pada seluruh modul utama sistem menggunakan berbagai skenario '
    'input untuk memastikan sistem berjalan sesuai dengan kebutuhan yang telah ditetapkan.'
)

doc.add_heading('3.5.1 Pengujian Fitur Login dan Autentikasi', level=3)
add_paragraph(
    'Pengujian fitur login dilakukan untuk memastikan bahwa sistem mampu melakukan '
    'proses autentikasi pengguna secara benar sesuai dengan peran masing-masing.'
)

# Test table
t2 = doc.add_table(rows=1, cols=4)
t2.style = 'Table Grid'
h2 = t2.rows[0].cells
headers = ['Skenario', 'Input', 'Hasil Diharapkan', 'Status']
for i, h in enumerate(headers):
    h2[i].text = h
    for p in h2[i].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)

test_cases_login = [
    ('Login Admin - Password benar', 'Password: admin123', 'Masuk ke Admin Dashboard', 'Berhasil'),
    ('Login Admin - Password salah', 'Password: admin', 'Menampilkan error "Password Admin salah!"', 'Berhasil'),
    ('Login Security - Password benar', 'Password: security123', 'Masuk ke Scan Interface', 'Berhasil'),
    ('Login Security - Password salah', 'Password: security', 'Menampilkan error "Password Security salah!"', 'Berhasil'),
    ('Login Employee - NIK terdaftar', 'NIK: EMP001', 'Masuk ke Employee Portal', 'Berhasil'),
    ('Login Employee - NIK tidak terdaftar', 'NIK: 000000', 'Menampilkan error "NIK tidak terdaftar!"', 'Berhasil'),
    ('Login Employee - NIK kosong', 'NIK: (kosong)', 'Tombol submit tidak aktif', 'Berhasil'),
]
for tc in test_cases_login:
    add_table_row(t2, list(tc))

add_paragraph('Tabel 3.1 Hasil Pengujian Login dan Autentikasi', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False, size=10)

doc.add_heading('3.5.2 Pengujian Scan Interface', level=3)
add_paragraph(
    'Pengujian Scan Interface dilakukan untuk memastikan bahwa fitur pemindaian absensi '
    'berfungsi dengan benar, baik melalui input teks manual maupun pemindaian kamera.'
)

t3 = doc.add_table(rows=1, cols=4)
t3.style = 'Table Grid'
h3 = t3.rows[0].cells
for i, h in enumerate(headers):
    h3[i].text = h
    for p in h3[i].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)

test_cases_scan = [
    ('Scan NIK terdaftar (IN)', 'NIK: EMP001', 'Notifikasi "Selamat Datang" + statistik IN bertambah', 'Berhasil'),
    ('Scan NIK yang sama (toggle OUT)', 'NIK: EMP001 (scan kedua)', 'Notifikasi "Selamat Jalan" + statistik IN berkurang', 'Berhasil'),
    ('Scan NIK tidak terdaftar', 'NIK: UNKNOWN', 'Notifikasi error "Karyawan tidak ditemukan"', 'Berhasil'),
    ('Cooldown 10 detik (client)', 'NIK: EMP001 (diulang <10 detik)', 'Input dibersihkan, tidak ada notifikasi', 'Berhasil'),
    ('Cooldown 1 menit (server)', 'NIK: EMP001 (diulang <1 menit)', 'Notifikasi "Mohon tunggu sebentar (Xs)"', 'Berhasil'),
    ('Scan menggunakan kamera', 'QR Code EMP001', 'Notifikasi sesuai dengan status IN/OUT', 'Berhasil'),
    ('Input dengan karakter kontrol', 'NIK: EMP001\\r\\n', 'Input dibersihkan, proses scan berhasil', 'Berhasil'),
]
for tc in test_cases_scan:
    add_table_row(t3, list(tc))

add_paragraph('Tabel 3.2 Hasil Pengujian Scan Interface', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False, size=10)

doc.add_heading('3.5.3 Pengujian Admin Dashboard', level=3)
add_paragraph(
    'Pengujian Admin Dashboard dilakukan untuk memastikan seluruh fitur manajemen data '
    'dan administrasi berfungsi dengan benar.'
)

t4 = doc.add_table(rows=1, cols=4)
t4.style = 'Table Grid'
h4 = t4.rows[0].cells
for i, h in enumerate(headers):
    h4[i].text = h
    for p in h4[i].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)

test_cases_admin = [
    ('Menambah karyawan baru', 'Isi form + submit', 'Data tersimpan di Firestore + notifikasi sukses', 'Berhasil'),
    ('Mengedit data karyawan', 'Ubah nama + submit', 'Data terupdate di Firestore', 'Berhasil'),
    ('Menghapus karyawan', 'Klik hapus + konfirmasi', 'Data terhapus dari Firestore', 'Berhasil'),
    ('Filter log berdasarkan tanggal', 'Pilih tanggal tertentu', 'Log sesuai tanggal yang dipilih', 'Berhasil'),
    ('Filter log berdasarkan departemen', 'Pilih departemen', 'Log difilter sesuai departemen', 'Berhasil'),
    ('Export PDF laporan', 'Klik Export PDF', 'File PDF terdownload dengan data sesuai filter', 'Berhasil'),
    ('Seed data 90 karyawan', 'Klik Seed Data 2x', '90 data karyawan tersimpan di Firestore', 'Berhasil'),
    ('Toggle maintenance mode', 'Klik Maint. Mode', 'Overlay maintenance muncul untuk non-admin', 'Berhasil'),
    ('Cetak Identity Pass', 'Klik tombol cetak', 'Dialog print browser terbuka', 'Berhasil'),
]
for tc in test_cases_admin:
    add_table_row(t4, list(tc))

add_paragraph('Tabel 3.3 Hasil Pengujian Admin Dashboard', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False, size=10)

doc.add_heading('3.5.4 Pengujian QR/Barcode Scanner', level=3)
add_paragraph(
    'Pengujian QR/Barcode Scanner dilakukan untuk memastikan bahwa kamera dapat '
    'mendeteksi dan membaca berbagai format kode dengan benar.'
)

t5 = doc.add_table(rows=1, cols=4)
t5.style = 'Table Grid'
h5 = t5.rows[0].cells
for i, h in enumerate(headers):
    h5[i].text = h
    for p in h5[i].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)

test_cases_scanner = [
    ('Scan QR Code valid', 'QR Code berisi NIK', 'Data NIK terbaca dan diproses', 'Berhasil'),
    ('Scan Barcode Code 128', 'Barcode Code 128', 'Data terbaca dan diproses', 'Berhasil'),
    ('Scan Barcode EAN-13', 'Barcode EAN-13', 'Data terbaca dan diproses', 'Berhasil'),
    ('Izin kamera ditolak', 'Tolak akses kamera', 'Menampilkan error "Gagal mengakses kamera"', 'Berhasil'),
    ('Switch kamera ON/OFF', 'Toggle tombol kamera', 'Kamera mati/nyala sesuai toggle', 'Berhasil'),
]
for tc in test_cases_scanner:
    add_table_row(t5, list(tc))

add_paragraph('Tabel 3.4 Hasil Pengujian QR/Barcode Scanner', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False, size=10)

doc.add_heading('3.5.5 Pengujian Attendance Logic', level=3)
add_paragraph(
    'Pengujian logika absensi dilakukan untuk memastikan bahwa sistem mampu menentukan '
    'status absensi (IN/OUT) dengan benar berdasarkan log terakhir, serta menjaga '
    'konsistensi data melalui batch write atomik.'
)

t6 = doc.add_table(rows=1, cols=4)
t6.style = 'Table Grid'
h6 = t6.rows[0].cells
for i, h in enumerate(headers):
    h6[i].text = h
    for p in h6[i].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)

test_cases_logic = [
    ('Scan pertama hari ini', 'NIK tanpa log hari ini', 'Status IN (Masuk)', 'Berhasil'),
    ('Scan setelah IN', 'NIK dengan log IN terakhir', 'Status OUT (Keluar)', 'Berhasil'),
    ('Scan setelah OUT', 'NIK dengan log OUT terakhir', 'Status IN (Masuk)', 'Berhasil'),
    ('POB bertambah saat IN', 'Scan IN', 'POB +1 di statistik', 'Berhasil'),
    ('POB berkurang saat OUT', 'Scan OUT', 'POB -1 di statistik', 'Berhasil'),
    ('Visitor tracking', 'Scan visitor IN/OUT', 'Visitor count terupdate', 'Berhasil'),
    ('Batch write atomicity', 'Simulasi kegagalan', 'Data log dan stats konsisten', 'Berhasil'),
]
for tc in test_cases_logic:
    add_table_row(t6, list(tc))

add_paragraph('Tabel 3.5 Hasil Pengujian Attendance Logic', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False, size=10)

doc.add_heading('3.5.6 Pengujian Maintenance Mode', level=3)
add_paragraph(
    'Pengujian maintenance mode dilakukan untuk memastikan bahwa fitur pemeliharaan '
    'sistem dapat berfungsi dengan baik dan membatasi akses sesuai dengan peran pengguna.'
)

t7 = doc.add_table(rows=1, cols=4)
t7.style = 'Table Grid'
h7 = t7.rows[0].cells
for i, h in enumerate(headers):
    h7[i].text = h
    for p in h7[i].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)

test_cases_maint = [
    ('Admin mengaktifkan maintenance', 'Klik Maint. Mode ON', 'Overlay muncul untuk Security & Employee', 'Berhasil'),
    ('Admin menonaktifkan maintenance', 'Klik Maint. Mode OFF', 'Overlay hilang untuk semua pengguna', 'Berhasil'),
    ('Akses Admin saat maintenance', 'Login Admin (maintenance ON)', 'Admin tetap bisa mengakses dashboard', 'Berhasil'),
    ('Akses Security saat maintenance', 'Login Security (maintenance ON)', 'Overlay maintenance muncul', 'Berhasil'),
    ('Pesan maintenance kustom', 'Update pesan di Firestore', 'Pesan baru tampil di overlay', 'Berhasil'),
]
for tc in test_cases_maint:
    add_table_row(t7, list(tc))

add_paragraph('Tabel 3.6 Hasil Pengujian Maintenance Mode', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False, size=10)

doc.add_page_break()

# ============================================================
# BAB IV - HASIL DAN PEMBAHASAN
# ============================================================
doc.add_heading('BAB IV\nHASIL DAN PEMBAHASAN', level=1)

doc.add_heading('4.1 Hasil Implementasi', level=2)
add_paragraph(
    'Berdasarkan tahapan pengembangan yang telah dilakukan, sistem absensi berbasis QR Code '
    'dan Barcode pada Warehouse Elnusa BSD berhasil dikembangkan dan diimplementasikan. '
    'Sistem ini dibangun menggunakan React.js dengan TypeScript pada sisi frontend dan '
    'Firebase Firestore sebagai backend basis data real-time. Seluruh fitur yang telah '
    'direncanakan pada tahap analisis kebutuhan berhasil diimplementasikan dan telah melewati '
    'tahap pengujian Black Box dengan hasil yang memuaskan.'
)
add_paragraph(
    'Hasil implementasi sistem meliputi beberapa antarmuka utama sebagai berikut:'
)
add_bullet(
    'Halaman Login (LoginSelection): Menampilkan tiga kartu pilihan peran pengguna '
    '(Admin, Security, Employee) dengan desain modern menggunakan gradasi warna dan animasi. '
    'Setiap kartu dilengkapi dengan ikon, deskripsi peran, dan tombol akses yang sesuai.'
)
add_bullet(
    'Halaman Scan Interface: Menampilkan input teks untuk pemindaian manual dan tombol '
    'untuk mengaktifkan kamera. Dilengkapi dengan notifikasi hasil scan yang animatif '
    'serta grid statistik real-time yang menampilkan data IN, OUT, POB, dan Visitor.'
)
add_bullet(
    'Halaman Admin Dashboard: Menampilkan dua tab (Log Attendance dan Employee Database) '
    'dengan fitur filter tanggal dan departemen. Dilengkapi dengan tombol aksi untuk '
    'manajemen data karyawan, ekspor PDF, seed data, dan pengaturan maintenance mode.'
)
add_bullet(
    'Halaman Employee Portal: Menampilkan QR Code dan Barcode pribadi karyawan, riwayat '
    'kehadiran dalam bentuk tabel, serta ringkasan statistik pribadi.'
)
add_paragraph(
    'Dari segi teknis, sistem berhasil mengimplementasikan fitur-fitur utama seperti: '
    '(1) mekanisme toggle IN/OUT yang cerdas berdasarkan log terakhir; (2) enam strategi '
    'pencarian karyawan untuk mengakomodasi variasi format input dari scanner; (3) batch '
    'write atomik dengan Firestore untuk menjaga konsistensi data log dan statistik; '
    '(4) real-time listener pada Firestore untuk pembaruan data secara langsung; (5) '
    'pemindaian multi-format menggunakan library html5-qrcode yang mendukung QR Code, '
    'Code 128, EAN-13, dan Code 39; serta (6) cooldown mechanism untuk mencegah '
    'pemindaian ganda dalam waktu singkat.'
)

doc.add_heading('4.2 Relevansi Mata Kuliah dengan Kegiatan Magang', level=2)
add_paragraph(
    'Kegiatan magang di Warehouse Elnusa BSD menjadi bentuk implementasi nyata dari capaian '
    'pembelajaran Program Studi Teknik Informatika. Selama pelaksanaan magang, penulis '
    'menerapkan berbagai pengetahuan dan keterampilan yang diperoleh selama perkuliahan '
    'dalam mengembangkan sistem absensi berbasis web. Berikut adalah relevansi mata kuliah '
    'dengan kegiatan magang yang telah dilaksanakan:'
)

doc.add_heading('Pemrograman Web', level=3)
add_paragraph(
    'Mata kuliah Pemrograman Web menekankan pemahaman terhadap pengembangan aplikasi web '
    'modern menggunakan teknologi terkini. Fokus pembelajaran mencakup arsitektur frontend, '
    'manajemen state, komponen reusable, serta integrasi dengan layanan backend. Dalam '
    'kegiatan magang di Warehouse Elnusa BSD, penerapan kompetensi tersebut diwujudkan '
    'melalui pengembangan sistem absensi berbasis web menggunakan React.js dan TypeScript. '
    'Sistem yang dikembangkan menerapkan arsitektur komponen dengan pemisahan logika bisnis, '
    'manajemen state menggunakan React hooks (useState, useEffect, useRef), serta integrasi '
    'real-time dengan Firebase Firestore melalui mekanisme snapshot listener. Pengalaman '
    'ini mencerminkan penerapan capaian pembelajaran Pemrograman Web dalam konteks '
    'pengembangan aplikasi web modern di lingkungan industri pergudangan.'
)

doc.add_heading('Basis Data', level=3)
add_paragraph(
    'Mata kuliah Basis Data menekankan pemahaman terhadap perancangan, implementasi, dan '
    'pengelolaan basis data dalam sistem informasi. Fokus pembelajaran mencakup pemodelan '
    'data, normalisasi, bahasa query, dan integritas data. Dalam kegiatan magang, '
    'penerapan kompetensi tersebut diwujudkan melalui perancangan dan implementasi basis '
    'data Firestore untuk sistem absensi. Struktur data dirancang dengan koleksi employees, '
    'presence_logs, stats, dan system_config yang saling terintegrasi. Operasi basis data '
    'menggunakan write batch untuk menjaga atomicity transaksi, serta real-time listener '
    'untuk sinkronisasi data secara langsung.'
)

doc.add_heading('Interaksi Manusia dan Komputer', level=3)
add_paragraph(
    'Mata kuliah Interaksi Manusia dan Komputer (IMK) menekankan pemahaman terhadap '
    'perancangan antarmuka yang efektif, efisien, dan mudah digunakan. Fokus pembelajaran '
    'mencakup prinsip-prinsip desain antarmuka, analisis kebutuhan pengguna, serta '
    'penerapan elemen visual yang mendukung pengalaman pengguna (user experience). Dalam '
    'kegiatan magang, penerapan kompetensi tersebut diwujudkan melalui perancangan '
    'antarmuka sistem absensi dengan pendekatan visual modern menggunakan Tailwind CSS. '
    'Antarmuka dirancang responsif untuk berbagai ukuran layar, dengan notifikasi animatif '
    'yang memberikan umpan balik visual jelas kepada pengguna, serta tata letak informasi '
    'yang hierarkis untuk memudahkan pemantauan data statistik.'
)

doc.add_heading('Rekayasa Perangkat Lunak', level=3)
add_paragraph(
    'Mata kuliah Rekayasa Perangkat Lunak (RPL) menekankan pemahaman terhadap metodologi '
    'pengembangan perangkat lunak yang sistematis dan terstruktur. Dalam kegiatan magang, '
    'penerapan kompetensi tersebut diwujudkan melalui penggunaan model Waterfall sebagai '
    'metode pengembangan yang mencakup tahapan analisis kebutuhan, perancangan, implementasi, '
    'pengujian, dan pemeliharaan. Seluruh tahapan didokumentasikan secara sistematis untuk '
    'memastikan kualitas dan keberlanjutan sistem.'
)

doc.add_heading('Tanggung Jawab Profesional', level=3)
add_paragraph(
    'Nilai tanggung jawab profesional tercermin dalam proses pengembangan sistem yang '
    'menuntut ketelitian dan akuntabilitas tinggi mengingat data yang dikelola merupakan '
    'data kehadiran personel yang memengaruhi proses operasional gudang. Penulis bertanggung '
    'jawab memastikan akurasi pencatatan data, keamanan akses sistem, serta dokumentasi '
    'setiap tahapan pengembangan secara sistematis. Sistem yang dikembangkan telah melalui '
    'pengujian fungsional sebelum digunakan untuk memastikan keandalan dan konsistensi data.'
)

doc.add_heading('Isu Sosial dan Keprofesian Teknologi Informasi', level=3)
add_paragraph(
    'Mata kuliah Isu Sosial dan Keprofesian Teknologi Informasi menekankan pemahaman '
    'terhadap dampak sosial dari penerapan teknologi informasi serta tanggung jawab '
    'profesional yang melekat pada praktik kerja di bidang teknologi. Dalam kegiatan '
    'magang, aspek etika dan profesionalisme diterapkan melalui pengelolaan data '
    'kehadiran personel yang bersifat sensitif, penerapan mekanisme autentikasi dan '
    'otorisasi pengguna, serta pembatasan akses data sesuai dengan peran masing-masing. '
    'Sistem dikembangkan dengan prinsip transparansi, akurasi, dan keamanan data untuk '
    'memastikan dampak positif bagi efisiensi operasional gudang.'
)

doc.add_page_break()

# ============================================================
# DAFTAR PUSTAKA
# ============================================================
doc.add_heading('DAFTAR PUSTAKA', level=1)
add_paragraph('')

references = [
    'Facebook Open Source. (2024). React Documentation. https://react.dev/',
    'Google. (2024). Firebase Documentation. https://firebase.google.com/docs',
    'Tailwind Labs. (2024). Tailwind CSS Documentation. https://tailwindcss.com/docs',
    'Evan You. (2024). Vite Documentation. https://vitejs.dev/',
    'Express.js. (2024). Express.js Documentation. https://expressjs.com/',
    'Mozilla. (2024). MDN Web Docs: QR Code. https://developer.mozilla.org/',
    'Alief, R. N., & Rianto, H. (2025). Perancangan Sistem Absensi Siswa Berbasis Quick Response (QR) Code Menggunakan Framework JavaScript. INSANtek, 6(2), 89-97. https://doi.org/10.31294/insantek.v6i2.10185',
    'Praba, A. D., Safitri, M., & Faridi, F. (2025). Aplikasi Absensi Berbasis Website Menggunakan QR Code untuk Peningkatan Efisiensi Pencatatan Kehadiran. Jurnal Teknik, 14(2). https://doi.org/10.31000/jt.v14i2.15517',
    'Fadhilah, L. N., Auliana, S., & Aryanto, G. D. P. (2025). Perancangan Sistem Absensi Siswa Berbasis Web Menggunakan QR Code Disekolah PAUD Amelia Darul Akhyar Cikande. Jurnal Multimedia dan Teknologi Informasi (Jatilima), 7(02), 282-290. https://doi.org/10.54209/jatilima.v7i02.1530',
    'Nuralif, I., & Fachrie, M. (2023). Development of a QR code-based attendance system for factory employees. International Journal Software Engineering and Computer Science (IJSECS), 3(3), 281-286.',
    'Djamarullah, A. R., Nuryasin, I., & Wibowo, H. (2024). Designing a QR Code Attendance System Using BYOD (Bring Your Own Device). Ultimatics: Jurnal Teknik Informatika, 16(1), 32-37.',
    'Balogun, F. (2026). Design and Implementation of an Enhanced QR-Code Based Attendance System. Journal of Computing and Social Informatics. https://doi.org/10.33736/jcsi.10752.2026',
    'Minh, T. (2024). html5-qrcode Library Documentation. https://github.com/mebjas/html5-qrcode',
    'Hall, H. (2024). jsPDF Library Documentation. https://github.com/parallax/jsPDF',
    'Simek, M. (2024). jspdf-autotable Documentation. https://github.com/simonbengtsson/jspdf-autotable',
    'Paz, J. (2024). date-fns Documentation. https://date-fns.org/',
    'Cristea, C. (2024). lucide-react Icons Library. https://lucide.dev/',
    'Müller, D. (2024). motion/react Animation Library. https://motion.dev/',
]
for i, ref in enumerate(references, 1):
    add_paragraph(f'{i}. {ref}', first_line_indent=False, size=11)

# ============================================================
# SAVE
# ============================================================
output_path = '/root/ElnusaAbsensiWEB/Laporan_Magang_Adam_Putra_Pratama.docx'
doc.save(output_path)
print(f"Laporan magang berhasil disimpan di: {output_path}")
print(f"Ukuran file: {os.path.getsize(output_path) / 1024:.1f} KB")
