#!/usr/bin/env python3
"""Generate LAPORAN AKHIR MAGANG - Template Sheetify untuk ElnusaAbsensiWEB."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ============================================================
# PAGE SETUP
# ============================================================
for section in doc.sections:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

# ============================================================
# STYLES
# ============================================================
style_normal = doc.styles['Normal']
font = style_normal.font
font.name = 'Times New Roman'
font.size = Pt(12)
style_normal.paragraph_format.line_spacing = 1.5
style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
style_normal.paragraph_format.space_after = Pt(6)

h1 = doc.styles['Heading 1']
h1.font.name = 'Times New Roman'
h1.font.size = Pt(14)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0, 0, 0)
h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
h1.paragraph_format.space_before = Pt(18)
h1.paragraph_format.space_after = Pt(12)

h2 = doc.styles['Heading 2']
h2.font.name = 'Times New Roman'
h2.font.size = Pt(12)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0, 0, 0)
h2.paragraph_format.space_before = Pt(12)
h2.paragraph_format.space_after = Pt(6)

h3 = doc.styles['Heading 3']
h3.font.name = 'Times New Roman'
h3.font.size = Pt(12)
h3.font.bold = True
h3.font.italic = True
h3.font.color.rgb = RGBColor(0, 0, 0)
h3.paragraph_format.space_before = Pt(6)
h3.paragraph_format.space_after = Pt(3)


def add_para(text, bold=False, italic=False, alignment=None, size=12, space_after=6, first_line_indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = alignment or WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    if first_line_indent and alignment is None:
        p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def add_numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F0F0F0')
    p.paragraph_format.element.get_or_add_pPr().append(shading)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p


def add_figure_caption(text, source=""):
    add_para(text, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=10, first_line_indent=False)
    if source:
        add_para(f"Sumber: {source}", italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=9, first_line_indent=False)


def add_image(image_path, width_cm=14):
    from docx.shared import Cm as DocxCm
    img_para = doc.add_paragraph()
    img_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.paragraph_format.space_before = Pt(6)
    img_para.paragraph_format.space_after = Pt(6)
    img_para.paragraph_format.line_spacing = 1.0
    run = img_para.add_run()
    run.add_picture(image_path, width=DocxCm(width_cm))


def add_table_row(table, cells, bold=False, size=10):
    row = table.add_row()
    for i, txt in enumerate(cells):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(txt))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.bold = bold
    return row


# ============================================================
# COVER
# ============================================================
for _ in range(3):
    add_para('', space_after=0)

add_para('LAPORAN AKHIR', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=18, first_line_indent=False)
add_para('MAGANG MBKM-MANDIRI', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=16, first_line_indent=False)
add_para('', space_after=0, first_line_indent=False)

add_para(
    'PENGEMBANGAN SISTEM ABSENSI BERBASIS QR CODE DAN BARCODE\n'
    'PADA WAREHOUSE ELNUSA BSD MENGGUNAKAN REACT.JS DAN FIREBASE',
    bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=14, first_line_indent=False
)
add_para('', space_after=0, first_line_indent=False)
add_para('', space_after=0, first_line_indent=False)

add_para('Diajukan Untuk Memenuhi Persyaratan MBKM-Mandiri', alignment=WD_ALIGN_PARAGRAPH.CENTER, size=12, first_line_indent=False)
add_para('Pada Fakultas Ilmu Komputer Universitas Esa Unggul', alignment=WD_ALIGN_PARAGRAPH.CENTER, size=12, first_line_indent=False)
add_para('', space_after=0, first_line_indent=False)

add_para('Disusun Oleh:', alignment=WD_ALIGN_PARAGRAPH.CENTER, size=12, first_line_indent=False)
add_para('Adam Putra Pratama', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=14, first_line_indent=False)
add_para('20230801402', alignment=WD_ALIGN_PARAGRAPH.CENTER, size=12, first_line_indent=False)
add_para('', space_after=0, first_line_indent=False)
add_para('', space_after=0, first_line_indent=False)
add_para('', space_after=0, first_line_indent=False)

add_para(
    'PROGRAM STUDI TEKNIK INFORMATIKA\n'
    'FAKULTAS ILMU KOMPUTER\n'
    'UNIVERSITAS ESA UNGGUL\n'
    'TANGERANG',
    bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=12, first_line_indent=False
)
add_para('', space_after=0, first_line_indent=False)
add_para('2026', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=12, first_line_indent=False)

doc.add_page_break()

# ============================================================
# LEMBAR PENGESAHAN
# ============================================================
add_para('LEMBAR PENGESAHAN', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=14, first_line_indent=False)
add_para('LAPORAN MAGANG', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=14, first_line_indent=False)
add_para('', space_after=0)

add_para(
    'Laporan magang ini diajukan oleh:',
    alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False
)
add_para('Nama            : Adam Putra Pratama', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False)
add_para('NIM             : 20230801402', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False)
add_para('Program Studi   : Teknik Informatika', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False)
add_para('Judul Laporan   : Pengembangan Sistem Absensi Berbasis QR Code dan Barcode pada Warehouse Elnusa BSD Menggunakan React.js dan Firebase', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False)

add_para('')
add_para('Telah berhasil dipertahankan di hadapan Dewan Penguji dan diterima sebagai bagian persyaratan yang diperlukan untuk memperoleh gelar Sarjana Komputer pada Program Studi Teknik Informatika, Fakultas Ilmu Komputer, Universitas Esa Unggul.', first_line_indent=True)
add_para('')

# Signatures section
add_para('Pembimbing Magang,', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False)
add_para('')
add_para('')
add_para('( .................................................. )', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False)
add_para('')
add_para('')
add_para('Dosen Pembimbing,', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False)
add_para('')
add_para('')
add_para('( .................................................. )', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False)
add_para('')
add_para('')
add_para('Mengetahui,', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False)
add_para('')
add_para('Kaprodi Teknik Informatika,', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False)
add_para('')
add_para('')
add_para('( .................................................. )', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False)

doc.add_page_break()

# ============================================================
# KATA PENGANTAR
# ============================================================
doc.add_heading('KATA PENGANTAR', level=1)
add_para('')
add_para(
    'Puji syukur kehadirat Tuhan Yang Maha Esa atas rahmat dan karunia-Nya sehingga Praktikan dapat '
    'menyusun dan menyelesaikan Laporan Akhir Magang MBKM-Mandiri yang berjudul "Pengembangan Sistem '
    'Absensi Berbasis QR Code dan Barcode pada Warehouse Elnusa BSD Menggunakan React.js dan Firebase" '
    'tepat pada waktunya.'
)
add_para(
    'Pelaksanaan Praktik Kerja Lapangan ini memberikan banyak pengalaman berharga serta wawasan baru '
    'yang sangat bermanfaat bagi pengembangan diri Praktikan, baik secara akademis maupun profesional. '
    'Praktikan menyadari bahwa keberhasilan penyusunan laporan ini tidak terlepas dari dukungan, '
    'bimbingan, dan bantuan dari berbagai pihak. Oleh karena itu, Praktikan ingin menyampaikan '
    'ucapan terima kasih yang sebesar-besarnya kepada:'
)
add_para(
    'Bapak Dr. Ir. Arief Kusuma A.P., S.T., M.B.A., IPU., ASEAN Eng., selaku Rektor Universitas Esa Unggul.\n'
    'Bapak Dr. Ir. Gerry Firmansyah, S.T. M.Kom, selaku Dekan Fakultas Ilmu Komputer.\n'
    'Ibu Dr. Riya Widayanti, S.Kom, MMSI, selaku Ketua Program Studi Teknik Informatika.\n'
    'Ibu Suryani, S.Si., M.Si., selaku Pembimbing Akademik.\n'
    '[Nama Dosen Pembimbing], selaku Dosen Pembimbing Magang.\n'
    '[Nama Pembimbing Lapangan], selaku Pembimbing Lapangan di Warehouse Elnusa BSD.\n'
    'Seluruh staf dan karyawan Warehouse Elnusa BSD yang telah membantu selama proses magang.\n'
    'Keluarga dan rekan-rekan yang selalu memberikan motivasi dan dukungan.'
)
add_para(
    'Praktikan menyadari bahwa laporan ini masih jauh dari sempurna. Oleh karena itu, kritik dan saran '
    'yang membangun sangat Praktikan harapkan untuk perbaikan di masa mendatang. Semoga laporan ini '
    'dapat memberikan manfaat bagi pengembangan ilmu pengetahuan dan teknologi, khususnya di bidang '
    'sistem informasi absensi digital.'
)
add_para('')
add_para('Tangerang, ... 2026', alignment=WD_ALIGN_PARAGRAPH.RIGHT, first_line_indent=False)
add_para('', space_after=0, first_line_indent=False)
add_para('Adam Putra Pratama', alignment=WD_ALIGN_PARAGRAPH.RIGHT, first_line_indent=False)

doc.add_page_break()

# ============================================================
# DAFTAR ISI (template)
# ============================================================
doc.add_heading('DAFTAR ISI', level=1)
add_para('')
toc_items = [
    ('KATA PENGANTAR', 'ii'),
    ('DAFTAR ISI', 'iii'),
    ('DAFTAR GAMBAR', 'iv'),
    ('DAFTAR TABEL', 'v'),
    ('DAFTAR LAMPIRAN', 'vi'),
    ('', ''),
    ('BAB 1 PENDAHULUAN', '1'),
    ('    1.1 Latar Belakang', '1'),
    ('    1.2 Maksud dan Tujuan Praktik Kerja Lapangan', '3'),
    ('    1.3 Manfaat Pelaksanaan Magang', '4'),
    ('    1.4 Pemilihan Instansi', '5'),
    ('    1.5 Waktu dan Pelaksanaan Tempat', '5'),
    ('    1.6 Kerangka Berpikir', '6'),
    ('', ''),
    ('BAB 2 LANDASAN TEORI', '7'),
    ('    2.1 React.js dan TypeScript', '7'),
    ('    2.2 Firebase Firestore', '8'),
    ('    2.3 QR Code dan Barcode', '9'),
    ('    2.4 Tailwind CSS', '10'),
    ('    2.5 Vite dan Express.js', '11'),
    ('    2.6 html5-qrcode', '11'),
    ('    2.7 jsPDF dan jspdf-autotable', '12'),
    ('    2.8 Autentikasi dan Otorisasi Pengguna', '12'),
    ('', ''),
    ('BAB 3 KEADAAN UMUM PERUSAHAAN', '13'),
    ('    3.1 Profil Perusahaan', '13'),
    ('    3.2 Visi dan Misi', '14'),
    ('    3.3 Layanan Instansi', '14'),
    ('    3.4 Struktur Organisasi', '15'),
    ('    3.5 Logo Instansi', '16'),
    ('', ''),
    ('BAB 4 ANALISIS HASIL DAN PEMBAHASAN', '17'),
    ('    4.1 Tools dan Teknologi Pengembangan', '17'),
    ('    4.2 Rancangan UML', '18'),
    ('    4.3 Rancangan DFD', '22'),
    ('    4.4 Rancangan ERD', '23'),
    ('    4.5 Proyek yang Dikerjakan', '24'),
    ('    4.6 User Interface Website', '35'),
    ('    4.7 Pemetaan dan Bukti Konversi Mata Kuliah', '40'),
    ('', ''),
    ('BAB 5 PENUTUP', '42'),
    ('    5.1 Kesimpulan', '42'),
    ('    5.2 Saran', '43'),
    ('', ''),
    ('DAFTAR PUSTAKA', '44'),
    ('LAMPIRAN', '46'),
]
for item, page in toc_items:
    if not item:
        add_para('', space_after=0, first_line_indent=False)
    else:
        dots = '.' * max(3, 60 - len(item))
        add_para(f'{item} {dots} {page}', first_line_indent=False)

doc.add_page_break()

# ============================================================
# DAFTAR GAMBAR
# ============================================================
doc.add_heading('DAFTAR GAMBAR', level=1)
add_para('')
gambar_list = [
    'Gambar 1.1 Waktu dan Pelaksanaan Tempat',
    'Gambar 2.1 React.js Logo',
    'Gambar 2.2 Firebase Logo',
    'Gambar 2.3 QR Code',
    'Gambar 2.4 Tailwind CSS',
    'Gambar 2.5 JavaScript',
    'Gambar 2.6 html5-qrcode',
    'Gambar 2.7 jsPDF',
    'Gambar 3.1 Layanan Warehouse Elnusa BSD',
    'Gambar 3.2 Struktur Organisasi Warehouse Elnusa BSD',
    'Gambar 3.3 Logo Elnusa',
    'Gambar 4.1 Flowchart Sistem',
    'Gambar 4.2 Use Case Diagram',
    'Gambar 4.3 Class Diagram',
    'Gambar 4.4 Activity Diagram',
    'Gambar 4.5 Sequence Diagram',
    'Gambar 4.6 DFD Level 0',
    'Gambar 4.7 DFD Level 1',
    'Gambar 4.8 Rancangan ERD',
    'Gambar 4.9 Halaman Login (LoginSelection)',
    'Gambar 4.10 Halaman Scan Interface',
    'Gambar 4.11 Halaman Admin Dashboard',
    'Gambar 4.12 Halaman Employee Portal',
    'Gambar 4.13 Scanner Kamera',
    'Gambar 4.14 Icon Aplikasi PWA (180x180 pixel)',
    'Gambar 4.15 Icon Aplikasi PWA Ukuran 120x120 pixel',
    'Gambar 4.16 Tampilan WarehouseDashboard POB Monitoring',
    'Gambar 4.17 Tampilan PWA Terinstal di Layar Utama iPhone',
]
for i, g in enumerate(gambar_list, 1):
    add_para(g, first_line_indent=False)

doc.add_page_break()

# ============================================================
# DAFTAR TABEL
# ============================================================
doc.add_heading('DAFTAR TABEL', level=1)
add_para('')
tabel_list = [
    'Tabel 1.1 Pemilihan Instansi',
    'Tabel 4.1 Tools dan Teknologi Pengembangan',
    'Tabel 4.2 Pemetaan dan Bukti (Evidence) Konversi Mata Kuliah',
]
for t in tabel_list:
    add_para(t, first_line_indent=False)

doc.add_page_break()

# ============================================================
# DAFTAR LAMPIRAN
# ============================================================
doc.add_heading('DAFTAR LAMPIRAN', level=1)
add_para('')
lampiran_list = [
    'Lampiran 1 Surat Permohonan Magang',
    'Lampiran 2 Surat Penerimaan Magang',
    'Lampiran 3 Learning Agreement',
    'Lampiran 4 Sertifikat dan Penilaian Magang',
    'Lampiran 5 Dokumentasi Magang',
    'Lampiran 6 Riwayat Hidup',
]
for l in lampiran_list:
    add_para(l, first_line_indent=False)

doc.add_page_break()

# ============================================================
# BAB 1 - PENDAHULUAN
# ============================================================
doc.add_heading('BAB 1\nPENDAHULUAN', level=1)

doc.add_heading('1.1 Latar Belakang', level=2)
add_para(
    'Perkembangan teknologi informasi dan komunikasi telah membawa perubahan signifikan dalam '
    'berbagai aspek kehidupan, termasuk di dunia industri dan pergudangan. Teknologi digital '
    'telah menjadi kebutuhan pokok bagi perusahaan untuk meningkatkan efisiensi operasional, '
    'akurasi data, serta kecepatan dalam pengambilan keputusan. Salah satu aspek operasional '
    'yang krusial dalam manajemen pergudangan adalah sistem absensi atau pencatatan kehadiran '
    'karyawan dan tamu yang masuk dan keluar dari area gudang.'
)
add_para(
    'Warehouse Elnusa BSD merupakan unit pergudangan dari PT Elnusa Tbk yang bergerak di bidang '
    'logistik dan penyimpanan. Dalam operasional sehari-harinya, Warehouse Elnusa BSD memerlukan '
    'sistem pencatatan kehadiran yang akurat dan efisien untuk memantau personel yang berada di '
    'dalam area gudang. Personel yang dimaksud mencakup karyawan tetap, karyawan tidak tetap, '
    'serta tamu atau visitor yang memiliki keperluan di area gudang.'
)
add_para(
    'Berdasarkan observasi awal yang dilakukan praktikan, sistem absensi yang berjalan di Warehouse '
    'Elnusa BSD masih menggunakan metode manual, yaitu pencatatan melalui buku tamu dan formulir '
    'kertas. Metode manual ini memiliki beberapa kelemahan, antara lain: (1) antrean panjang pada '
    'saat jam masuk dan keluar kerja karena proses pencatatan yang lambat; (2) rawan terjadinya '
    'kesalahan pencatatan data akibat tulisan tangan yang tidak jelas; (3) kesulitan dalam '
    'merekap data kehadiran secara real-time karena data tersebar di buku catatan fisik; '
    '(4) potensi manipulasi data kehadiran karena tidak ada mekanisme verifikasi yang ketat; '
    'serta (5) keterbatasan dalam memantau jumlah personel yang sedang berada di dalam area gudang '
    '(Person On Board/POB) secara langsung.'
)
add_para(
    'Untuk mengatasi permasalahan tersebut, diperlukan sebuah sistem absensi digital yang mampu '
    'melakukan pencatatan kehadiran secara otomatis, akurat, dan real-time. Teknologi QR Code '
    'dan Barcode menjadi pilihan yang tepat karena keduanya dapat menyimpan data identitas '
    'unik setiap personel dalam bentuk kode dua dimensi maupun satu dimensi yang mudah dibaca '
    'oleh pemindai optik. Setiap karyawan akan memiliki identitas digital berupa QR Code dan '
    'Barcode yang dapat dipindai oleh petugas security di pintu gerbang masuk maupun keluar '
    'gudang.'
)
add_para(
    'Proyek pengembangan sistem absensi ini dibangun menggunakan teknologi web modern, yaitu '
    'React.js dengan TypeScript untuk antarmuka pengguna (frontend) dan Firebase sebagai layanan '
    'backend yang menyediakan basis data real-time (Firestore). Pemilihan React.js didasarkan '
    'pada kemampuannya dalam membangun antarmuka pengguna yang interaktif, responsif, dan '
    'reaktif terhadap perubahan data secara real-time. Sementara itu, Firebase Firestore dipilih '
    'karena menyediakan layanan basis data NoSQL yang terkelola sepenuhnya dengan kemampuan '
    'sinkronisasi data secara real-time, sehingga cocok untuk aplikasi yang membutuhkan '
    'pembaruan data secara langsung seperti sistem absensi.'
)

doc.add_heading('1.2 Maksud dan Tujuan Praktik Kerja Lapangan', level=2)
add_para('Adapun maksud dilaksanakannya kegiatan MBKM Mandiri antara lain:', first_line_indent=False)
add_numbered('Memenuhi persyaratan kurikulum Sarjana Strata-1 Program Studi Teknik Informatika.')
add_numbered('Menerapkan ilmu yang telah diperoleh selama di perkuliahan dengan mengikuti program MBKM Mandiri.')
add_numbered('Menambah wawasan mengenai suatu bidang pekerjaan di industri pergudangan dan logistik.')

add_para('Adapun tujuan dilaksanakannya kegiatan MBKM Mandiri antara lain:', first_line_indent=False)
add_numbered('Mempersiapkan mahasiswa sebagai Sumber Daya Manusia (SDM) yang berkualitas dalam menghadapi persaingan dunia kerja.')
add_numbered('Meningkatkan pengetahuan serta keterampilan yang sesuai dengan latar belakang bidang studi Teknik Informatika.')
add_numbered('Menerapkan ilmu akademis yang telah diperoleh selama di bangku perkuliahan, khususnya dalam pengembangan sistem berbasis web.')
add_numbered('Memberikan kontribusi kepada perusahaan yang sesuai dengan ilmu pengetahuan yang ditekuni praktikan.')

doc.add_heading('1.3 Manfaat Pelaksanaan Magang', level=2)

doc.add_heading('Bagi Mahasiswa', level=3)
add_bullet('Meningkatkan pengetahuan, keterampilan, dan pengalaman praktis dalam pengembangan aplikasi web menggunakan React.js dan Firebase.')
add_bullet('Mengaplikasikan ilmu yang diperoleh selama perkuliahan ke dalam penyelesaian permasalahan nyata, khususnya dalam pengembangan sistem absensi digital.')
add_bullet('Mengembangkan kemampuan profesional, komunikasi, kerja sama tim, dan pemecahan masalah sebagai bekal menghadapi dunia kerja.')

doc.add_heading('Bagi Instansi', level=3)
add_bullet('Membantu pengembangan sistem absensi digital sebagai solusi untuk menggantikan pencatatan kehadiran manual.')
add_bullet('Mendukung peningkatan efektivitas pengelolaan data kehadiran melalui sistem yang terintegrasi dan real-time.')
add_bullet('Memberikan kontribusi berupa penerapan teknologi informasi yang dapat mendukung transformasi digital di lingkungan Warehouse Elnusa BSD.')

doc.add_heading('Bagi Universitas', level=3)
add_bullet('Memperkuat kerja sama antara Universitas Esa Unggul dengan industri dalam pelaksanaan program MBKM Mandiri.')
add_bullet('Menjadi bahan evaluasi terhadap kesesuaian kurikulum dengan kebutuhan dunia kerja, khususnya pada bidang pengembangan sistem informasi.')
add_bullet('Meningkatkan kualitas lulusan melalui pengalaman belajar berbasis praktik sehingga menghasilkan sumber daya manusia yang kompeten.')

doc.add_heading('1.4 Pemilihan Instansi', level=2)
add_para(
    'Praktikan melaksanakan kegiatan MBKM Mandiri di Warehouse Elnusa BSD yang merupakan unit '
    'pergudangan dari PT Elnusa Tbk. Pemilihan instansi ini didasarkan pada relevansi bidang '
    'kerja dengan program studi Teknik Informatika, khususnya dalam pengembangan sistem informasi '
    'berbasis web.'
)
t1 = doc.add_table(rows=1, cols=3)
t1.style = 'Table Grid'
for i, h in enumerate(['Aspek', 'Keterangan']):
    t1.rows[0].cells[i].text = h
    for r in t1.rows[0].cells[i].paragraphs:
        for run in r.runs:
            run.bold = True
            run.font.size = Pt(10)
add_table_row(t1, ['Nama Instansi', 'Warehouse Elnusa BSD (PT Elnusa Tbk)'])
add_table_row(t1, ['Bidang Usaha', 'Logistik dan Pergudangan'])
add_table_row(t1, ['Lokasi', 'BSD, Tangerang'])
add_table_row(t1, ['Bagian Penempatan', 'IT & Sistem Informasi'])
add_table_row(t1, ['Waktu Pelaksanaan', '... s.d. ... 2026'])
add_table_row(t1, ['Proyek', 'Pengembangan Sistem Absensi Berbasis QR Code dan Barcode'])
add_figure_caption('Tabel 1.1 Pemilihan Instansi')

doc.add_heading('1.5 Waktu dan Pelaksanaan Tempat', level=2)
add_para(
    'Magang di Warehouse Elnusa BSD dilaksanakan selama ... bulan. Terhitung dari tanggal ... '
    'sampai dengan ... 2026. Kegiatan magang dilaksanakan secara luring (offline) di kantor '
    'Warehouse Elnusa BSD yang berlokasi di kawasan BSD, Tangerang. Jam kerja mengikuti '
    'ketentuan perusahaan, yaitu hari Senin sampai dengan Jumat pukul 08.00 hingga 17.00 WIB.'
)

add_figure_caption('[Gambar 1.1 Waktu dan Pelaksanaan Tempat]')

doc.add_heading('1.6 Kerangka Berpikir', level=2)
add_para(
    'Kerangka berpikir dalam laporan magang ini disusun secara sistematis untuk menggambarkan '
    'alur logis penulisan, mulai dari identifikasi masalah hingga penarikan kesimpulan.'
)

add_para('BAB 1 PENDAHULUAN', bold=True, first_line_indent=False)
add_para(
    'Bab ini menguraikan latar belakang pelaksanaan magang di Warehouse Elnusa BSD, maksud dan '
    'tujuan, manfaat bagi mahasiswa, instansi, dan universitas, pemilihan instansi, serta waktu '
    'dan tempat pelaksanaan magang.'
)
add_para('BAB 2 LANDASAN TEORI', bold=True, first_line_indent=False)
add_para(
    'Bab ini membahas landasan teori yang menjadi dasar ilmiah dalam perancangan dan pengembangan '
    'sistem absensi, meliputi konsep React.js, Firebase Firestore, QR Code, Barcode, Tailwind CSS, '
    'serta teknologi pendukung lainnya.'
)
add_para('BAB 3 KEADAAN UMUM PERUSAHAAN', bold=True, first_line_indent=False)
add_para(
    'Bab ini menjelaskan keadaan umum Warehouse Elnusa BSD sebagai tempat pelaksanaan magang, '
    'mencakup profil perusahaan, visi dan misi, layanan, struktur organisasi, serta logo instansi.'
)
add_para('BAB 4 ANALISIS HASIL DAN PEMBAHASAN', bold=True, first_line_indent=False)
add_para(
    'Bab ini memaparkan hasil analisis dan pembahasan terhadap proyek yang dikerjakan selama '
    'magang, mencakup tools dan teknologi, rancangan UML, DFD, ERD, implementasi proyek, serta '
    'tampilan antarmuka sistem.'
)
add_para('BAB 5 PENUTUP', bold=True, first_line_indent=False)
add_para(
    'Bab ini berisi kesimpulan yang merangkum keseluruhan hasil pelaksanaan magang dan pengembangan '
    'sistem absensi, serta saran untuk pengembangan selanjutnya.'
)

doc.add_page_break()

# ============================================================
# BAB 2 - LANDASAN TEORI
# ============================================================
doc.add_heading('BAB 2\nLANDASAN TEORI', level=1)

doc.add_heading('2.1 React.js dan TypeScript', level=2)
add_para(
    'React.js merupakan salah satu library JavaScript yang paling populer untuk membangun antarmuka '
    'pengguna (user interface) pada aplikasi web. Dikembangkan oleh Meta (sebelumnya Facebook), '
    'React.js menggunakan pendekatan komponen (component-based architecture) di mana setiap bagian '
    'dari antarmuka pengguna dibangun sebagai komponen yang independen, dapat digunakan kembali, '
    'dan memiliki state-nya sendiri. Pendekatan ini memudahkan pengembang dalam mengelola kompleksitas '
    'aplikasi web modern yang memiliki banyak interaksi dan pembaruan data secara dinamis (Facebook Open Source, 2024).'
)
add_para(
    'TypeScript adalah bahasa pemrograman yang merupakan superset dari JavaScript, yang menambahkan '
    'dukungan tipe data statis (static typing). Dengan adanya tipe data statis, TypeScript memungkinkan '
    'deteksi kesalahan pada tahap kompilasi (compile-time) sebelum kode dijalankan di browser. Pada proyek '
    'ini, React.js digunakan bersama dengan TypeScript untuk membangun seluruh komponen antarmuka sistem absensi.'
)
add_figure_caption('[Gambar 2.1 React.js]', 'https://react.dev/')

doc.add_heading('2.2 Firebase Firestore', level=2)
add_para(
    'Firebase adalah platform pengembangan aplikasi milik Google yang menyediakan berbagai layanan '
    'backend, termasuk basis data real-time (Firestore), autentikasi pengguna, penyimpanan file, '
    'dan hosting. Firebase Firestore adalah basis data NoSQL yang bersifat fleksibel, scalable, dan '
    'terkelola sepenuhnya di cloud. Data dalam Firestore disimpan dalam bentuk dokumen (document) '
    'yang terorganisir dalam koleksi (collection) (Google, 2024).'
)
add_para(
    'Salah satu keunggulan utama Firestore adalah kemampuannya untuk melakukan sinkronisasi data '
    'secara real-time menggunakan mekanisme snapshot listener. Dalam proyek ini, Firestore digunakan '
    'untuk menyimpan data karyawan (collection employees), log kehadiran (collection presence_logs), '
    'statistik harian (collection stats), serta konfigurasi sistem (collection system_config).'
)
add_figure_caption('[Gambar 2.2 Firebase]', 'https://firebase.google.com/')

doc.add_heading('2.3 QR Code dan Barcode', level=2)
add_para(
    'QR Code (Quick Response Code) adalah jenis kode matriks dua dimensi yang dapat menyimpan '
    'informasi dalam format teks, URL, atau data lainnya. QR Code dapat dibaca dengan cepat '
    'menggunakan kamera smartphone atau scanner khusus dari berbagai sudut pemindaian. Sementara itu, '
    'Barcode atau kode batang adalah representasi data optis yang dapat dibaca oleh mesin dalam '
    'bentuk garis-garis vertikal dengan ketebalan dan jarak yang bervariasi (Mozilla, 2024).'
)
add_para(
    'Dalam sistem absensi yang dikembangkan, setiap karyawan dan visitor memiliki identitas unik '
    'berupa NIK (Nomor Induk Karyawan) yang direpresentasikan dalam dua format kode, yaitu QR Code '
    'dan Barcode. Pada sisi pemindaian, sistem mendukung dua metode: input melalui scanner hardware '
    '(keyboard wedge) dan pemindaian melalui kamera perangkat.'
)
add_figure_caption('[Gambar 2.3 QR Code]')

doc.add_heading('2.4 Tailwind CSS', level=2)
add_para(
    'Tailwind CSS adalah framework CSS utility-first yang memungkinkan pengembang membangun '
    'antarmuka pengguna dengan cepat menggunakan kelas-kelas utilitas yang telah disediakan. '
    'Berbeda dengan framework CSS tradisional, Tailwind CSS menyediakan kelas-kelas utilitas '
    'tingkat rendah yang dapat dikombinasikan secara bebas (Tailwind Labs, 2024). Pada proyek ini, '
    'Tailwind CSS versi 4 digunakan untuk membangun seluruh antarmuka sistem absensi dengan desain '
    'yang responsif dan konsisten.'
)
add_figure_caption('[Gambar 2.4 Tailwind CSS]', 'https://tailwindcss.com/')

doc.add_heading('2.5 Vite dan Express.js', level=2)
add_para(
    'Vite adalah build tool modern untuk pengembangan aplikasi web yang menawarkan kecepatan '
    'pengembangan sangat tinggi dengan menggunakan ES Modules secara native (You, 2024). '
    'Express.js adalah framework web minimalis untuk Node.js yang digunakan untuk membangun server '
    'HTTP dan API (Express.js, 2024). Dalam proyek ini, Vite digunakan sebagai bundler untuk '
    'aplikasi React, sedangkan Express.js digunakan sebagai server pengembangan.'
)

doc.add_heading('2.6 html5-qrcode', level=2)
add_para(
    'html5-qrcode adalah library JavaScript ringan yang memungkinkan pemindaian QR Code dan '
    'Barcode langsung melalui kamera perangkat tanpa memerlukan plugin tambahan. Library ini '
    'mendukung berbagai format kode termasuk QR Code, Code 128, EAN-13, dan Code 39 (Minh, 2024). '
    'Dalam proyek ini, html5-qrcode digunakan pada komponen Scanner untuk memindai identitas '
    'karyawan melalui kamera.'
)
add_figure_caption('[Gambar 2.6 html5-qrcode]', 'https://github.com/mebjas/html5-qrcode')

doc.add_heading('2.7 jsPDF dan jspdf-autotable', level=2)
add_para(
    'jsPDF adalah library JavaScript untuk menghasilkan dokumen PDF secara client-side. '
    'Library ini memungkinkan pembuatan PDF dengan teks, gambar, dan elemen lainnya secara '
    'programatis (Hall, 2024). jspdf-autotable adalah plugin untuk jsPDF yang memudahkan '
    'pembuatan tabel dalam dokumen PDF dengan konfigurasi yang fleksibel (Simek, 2024). '
    'Dalam proyek ini, kedua library digunakan pada fitur Export PDF laporan kehadiran.'
)
add_figure_caption('[Gambar 2.7 jsPDF]', 'https://github.com/parallax/jsPDF')

doc.add_heading('2.8 Autentikasi dan Otorisasi Pengguna', level=2)
add_para(
    'Sistem absensi ini menerapkan tiga peran pengguna (role-based access control), yaitu Admin, '
    'Security, dan Employee. Admin memiliki akses penuh terhadap manajemen data dan pengaturan sistem. '
    'Security memiliki akses ke fitur pemindaian dan monitoring dashboard. Employee memiliki akses '
    'terbatas untuk melihat identitas digital dan riwayat kehadiran pribadi. Autentikasi untuk role '
    'Admin dan Security menggunakan password statis, sedangkan role Employee menggunakan verifikasi '
    'NIK terhadap data di Firestore.'
)

doc.add_page_break()

# ============================================================
# BAB 3 - KEADAAN UMUM PERUSAHAAN
# ============================================================
doc.add_heading('BAB 3\nKEADAAN UMUM PERUSAHAAN', level=1)

doc.add_heading('3.1 Profil Perusahaan', level=2)
add_para(
    'PT Elnusa Tbk merupakan perusahaan Badan Usaha Milik Negara (BUMN) yang bergerak di bidang '
    'jasa energi terintegrasi. Warehouse Elnusa BSD adalah unit pergudangan dari PT Elnusa Tbk '
    'yang berlokasi di kawasan BSD, Tangerang. Warehouse ini berfungsi sebagai pusat penyimpanan '
    'dan distribusi logistik perusahaan dengan standar operasional yang tinggi.'
)

doc.add_heading('3.2 Visi dan Misi', level=2)
add_para('Visi:', bold=True, first_line_indent=False)
add_para('"Menjadi perusahaan jasa energi terintegrasi yang unggul di tingkat regional."')
add_para('Misi:', bold=True, first_line_indent=False)
add_bullet('Memberikan solusi jasa energi yang terintegrasi dan bernilai tambah.')
add_bullet('Mengembangkan kompetensi inti dan sumber daya manusia yang unggul.')
add_bullet('Menerapkan tata kelola perusahaan yang baik (Good Corporate Governance).')
add_bullet('Berkontribusi pada pembangunan berkelanjutan dan lingkungan.')

doc.add_heading('3.3 Layanan Instansi', level=2)
add_para(
    'Warehouse Elnusa BSD menyediakan layanan pergudangan dan logistik yang mencakup penerimaan '
    'barang, penyimpanan, pengelolaan inventaris, dan distribusi. Dalam mendukung operasionalnya, '
    'warehouse memerlukan sistem informasi yang andal, termasuk sistem absensi untuk memantau '
    'kehadiran personel yang bekerja di area gudang.'
)
add_figure_caption('[Gambar 3.1 Layanan Warehouse Elnusa BSD]')

doc.add_heading('3.4 Struktur Organisasi', level=2)
add_para(
    'Struktur organisasi Warehouse Elnusa BSD dipimpin oleh Kepala Gudang yang membawahi beberapa '
    'divisi, antara lain Divisi Operasional, Divisi IT & Sistem Informasi, Divisi Keamanan (Security), '
    'dan Divisi Administrasi. Selama pelaksanaan magang, praktikan ditempatkan pada Divisi IT & '
    'Sistem Informasi yang bertanggung jawab atas pengembangan dan pemeliharaan sistem informasi di '
    'lingkungan warehouse.'
)
add_figure_caption('[Gambar 3.2 Struktur Organisasi Warehouse Elnusa BSD]')

doc.add_heading('3.5 Logo Instansi', level=2)
add_para(
    'PT Elnusa Tbk memiliki logo resmi yang mencerminkan identitas perusahaan sebagai penyedia '
    'jasa energi terintegrasi. Logo ini digunakan di seluruh dokumen resmi dan sistem informasi '
    'perusahaan.'
)
add_figure_caption('[Gambar 3.3 Logo Elnusa]')

doc.add_page_break()

# ============================================================
# BAB 4 - ANALISIS HASIL DAN PEMBAHASAN
# ============================================================
doc.add_heading('BAB 4\nANALISIS HASIL DAN PEMBAHASAN', level=1)

doc.add_heading('4.1 Tools dan Teknologi Pengembangan', level=2)
add_para(
    'Pengembangan sistem dalam penelitian ini didukung oleh beberapa tools yang dipilih sesuai '
    'kebutuhan, mulai dari penulisan kode, pengujian, hingga deployment aplikasi.'
)

t2 = doc.add_table(rows=1, cols=3)
t2.style = 'Table Grid'
for i, h in enumerate(['Tools', 'Fungsi', 'Jenis']):
    t2.rows[0].cells[i].text = h
    for r in t2.rows[0].cells[i].paragraphs:
        for run in r.runs:
            run.bold = True
            run.font.size = Pt(10)

tools_data = [
    ('Visual Studio Code', 'Editor kode untuk pengembangan aplikasi', 'Code Editor'),
    ('React.js + TypeScript', 'Library frontend untuk membangun antarmuka', 'Frontend'),
    ('Vite', 'Build tool dan bundler untuk aplikasi React', 'Build Tool'),
    ('Firebase Firestore', 'Basis data real-time (NoSQL)', 'Database'),
    ('Tailwind CSS v4', 'Framework CSS untuk styling', 'CSS Framework'),
    ('Express.js', 'Web server untuk development', 'Backend'),
    ('html5-qrcode', 'Library QR/Barcode scanner via kamera', 'Library'),
    ('react-qr-code', 'Library untuk generate QR Code', 'Library'),
    ('react-barcode', 'Library untuk generate Barcode', 'Library'),
    ('jsPDF + jspdf-autotable', 'Library untuk generate PDF', 'Library'),
    ('date-fns', 'Library untuk manipulasi tanggal', 'Library'),
    ('lucide-react', 'Library ikon untuk antarmuka', 'Library'),
    ('motion/react', 'Library animasi untuk React', 'Library'),
    ('Git', 'Version control system', 'Tools'),
]
for td in tools_data:
    add_table_row(t2, list(td))

add_figure_caption('Tabel 4.1 Tools dan Teknologi Pengembangan')

doc.add_heading('4.2 Rancangan UML', level=2)

doc.add_heading('4.2.1 Flowchart', level=3)
add_para(
    'Flowchart sistem absensi menggambarkan alur kerja sistem secara keseluruhan. Proses dimulai '
    'ketika pengguna membuka aplikasi dan memilih peran (role). Admin dan Security melakukan '
    'login menggunakan password, sedangkan Employee menggunakan NIK. Setelah login, masing-masing '
    'role diarahkan ke antarmuka yang sesuai dengan hak aksesnya. Security dapat melakukan scan '
    'NIK secara manual atau melalui kamera, Employee dapat melihat identitas digital dan riwayat, '
    'serta Admin dapat mengelola data dan pengaturan sistem.'
)
add_image('/root/ElnusaAbsensiWEB/diagrams/flowchart.png')
add_figure_caption('Gambar 4.1 Flowchart Sistem')

doc.add_heading('4.2.2 Use Case Diagram', level=3)
add_para(
    'Use case diagram menggambarkan interaksi antara tiga aktor (Admin, Security, Employee) '
    'dengan sistem absensi yang dikembangkan.'
)

# Use case list
add_para('Admin:', bold=True, first_line_indent=False)
add_bullet('Login Admin')
add_bullet('Mengelola Data Karyawan (Create, Read, Update, Delete)')
add_bullet('Mengelola Data Visitor')
add_bullet('Mengekspor Laporan PDF')
add_bullet('Melakukan Seed Data (pengisian data dummy)')
add_bullet('Mengaktifkan/Menonaktifkan Maintenance Mode')
add_bullet('Mencetak Identity Pass')

add_para('Security:', bold=True, first_line_indent=False)
add_bullet('Login Security')
add_bullet('Memindai NIK (input manual)')
add_bullet('Memindai QR/Barcode (kamera)')
add_bullet('Melihat Statistik Real-time (IN, OUT, POB, Visitor)')

add_para('Employee:', bold=True, first_line_indent=False)
add_bullet('Login NIK')
add_bullet('Melihat QR Code dan Barcode Pribadi')
add_bullet('Melihat Riwayat Kehadiran')

add_image('/root/ElnusaAbsensiWEB/diagrams/usecase.png')
add_figure_caption('Gambar 4.2 Use Case Diagram')

doc.add_heading('4.2.3 Class Diagram', level=3)
add_para(
    'Class diagram sistem absensi menggambarkan struktur kelas-kelas utama yang digunakan dalam '
    'sistem. Kelas utama meliputi Employee, PresenceLog, DailyStats, UserRole, dan PresenceType. '
    'Hubungan antar kelas menunjukkan bahwa setiap Employee memiliki banyak PresenceLog, dan setiap '
    'PresenceLog memiliki satu tipe PresenceType (IN atau OUT). DailyStats menyimpan agregasi data '
    'harian untuk statistik dashboard.'
)
add_image('/root/ElnusaAbsensiWEB/diagrams/class_diagram.png')
add_figure_caption('Gambar 4.3 Class Diagram')

doc.add_heading('4.2.4 Activity Diagram', level=3)
add_para(
    'Activity diagram proses scan absensi menggambarkan alur aktivitas ketika Security melakukan '
    'pemindaian NIK. Proses dimulai dengan input NIK (manual atau kamera), kemudian sistem '
    'membersihkan input dari karakter kontrol. Sistem mencari data karyawan menggunakan enam '
    'strategi pencarian berurutan. Jika ditemukan, sistem menentukan jenis absensi (IN atau OUT) '
    'berdasarkan log terakhir, lalu mencatat log dan memperbarui statistik secara atomik '
    'menggunakan batch write Firestore.'
)
add_image('/root/ElnusaAbsensiWEB/diagrams/activity_diagram.png')
add_figure_caption('Gambar 4.4 Activity Diagram')

doc.add_heading('4.2.5 Sequence Diagram', level=3)
add_para(
    'Sequence diagram proses scan menggambarkan interaksi antara aktor Security, sistem frontend '
    '(React), dan backend (Firestore) dalam skenario scan absensi. Security memasukkan NIK, '
    'frontend mengirim permintaan pencarian ke Firestore, Firestore mengembalikan data karyawan, '
    'frontend menentukan jenis absensi, kemudian menulis log dan memperbarui statistik secara '
    'bersamaan.'
)
add_image('/root/ElnusaAbsensiWEB/diagrams/sequence_diagram.png')
add_figure_caption('Gambar 4.5 Sequence Diagram')

doc.add_heading('4.3 Rancangan DFD', level=2)

doc.add_heading('4.3.1 DFD Level 0 (Diagram Konteks)', level=3)
add_para(
    'DFD Level 0 sistem absensi menggambarkan interaksi antara tiga entitas eksternal (Admin, '
    'Security, Employee) dengan sistem. Admin memberikan data konfigurasi dan data karyawan, '
    'serta menerima laporan. Security memberikan data scan dan menerima notifikasi serta '
    'statistik. Employee memberikan NIK dan menerima identitas digital serta riwayat kehadiran.'
)
add_image('/root/ElnusaAbsensiWEB/diagrams/dfd_level0.png')
add_figure_caption('Gambar 4.6 DFD Level 0 (Diagram Konteks)')

doc.add_heading('4.3.2 DFD Level 1', level=3)
add_para(
    'DFD Level 1 memecah proses sistem menjadi sub-proses yang lebih detail, meliputi: '
    '(1) Proses Login dan Autentikasi, (2) Proses Scan Absensi, (3) Proses Manajemen Data Karyawan, '
    '(4) Proses Export Laporan, dan (5) Proses Monitoring Statistik. Masing-masing sub-proses '
    'berinteraksi dengan penyimpanan data (data store) seperti employees, presence_logs, stats, '
    'dan system_config.'
)
add_image('/root/ElnusaAbsensiWEB/diagrams/dfd_level1.png')
add_figure_caption('Gambar 4.7 DFD Level 1')

doc.add_heading('4.4 Rancangan ERD', level=2)
add_para(
    'Entity Relationship Diagram (ERD) sistem absensi menggambarkan hubungan antar entitas data '
    'dalam Firestore. Entitas utama meliputi:'
)
add_bullet('Employees: Menyimpan data karyawan dan visitor (id, name, nik, department, isVisitor)')
add_bullet('PresenceLogs: Menyimpan catatan kehadiran (employeeId, type [IN/OUT], timestamp, date)')
add_bullet('Stats: Menyimpan statistik harian (in, out, pob, visitorIn, visitorOut, totalVisits)')
add_bullet('SystemConfig: Menyimpan konfigurasi sistem (maintenanceMode, message)')
add_image('/root/ElnusaAbsensiWEB/diagrams/erd.png')
add_figure_caption('Gambar 4.8 Rancangan ERD')

doc.add_heading('4.5 Proyek yang Dikerjakan', level=2)
add_para(
    'Proyek utama yang dikerjakan selama magang adalah pengembangan Sistem Absensi Berbasis QR Code '
    'dan Barcode. Sistem ini dibangun menggunakan React.js dengan TypeScript pada sisi frontend dan '
    'Firebase Firestore sebagai basis data. Berikut adalah rincian implementasi per file komponen.'
)

doc.add_heading('4.5.1 LoginSelection.tsx', level=3)
add_para(
    'Halaman login merupakan gerbang utama sistem yang menampilkan tiga kartu pilihan peran: '
    'System Admin, Security Officer, dan Employee Portal. Admin menggunakan password "admin123", '
    'Security menggunakan password "security123", sedangkan Employee login menggunakan NIK yang '
    'terdaftar di database Firestore. Autentikasi Employee dilakukan dengan mencari dokumen '
    'berdasarkan NIK di koleksi employees.'
)

add_code(
    'const handleEmployeeLogin = async (e: React.FormEvent) => {\n'
    '  e.preventDefault();\n'
    '  if (!nik) return;\n'
    '  setIsLoading(true);\n'
    '  try {\n'
    '    const docRef = doc(db, "employees", nik);\n'
    '    const docSnap = await getDoc(docRef);\n'
    '    if (docSnap.exists()) {\n'
    '      const empData = { id: docSnap.id, ...docSnap.data() } as Employee;\n'
    '      onSelectRole(UserRole.EMPLOYEE, empData);\n'
    '    } else {\n'
    '      setError("NIK tidak terdaftar!");\n'
    '    }\n'
    '  } catch (err) {\n'
    '    setError("Gagal menghubungkan ke server.");\n'
    '  } finally {\n'
    '    setIsLoading(false);\n'
    '  }\n'
    '};'
)
add_figure_caption('[Gambar 4.9 Halaman Login (LoginSelection)]')

doc.add_heading('4.5.2 Scanner.tsx', level=3)
add_para(
    'Komponen Scanner menggunakan library html5-qrcode untuk memindai QR Code dan Barcode melalui '
    'kamera perangkat. Kamera yang digunakan adalah kamera belakang (facingMode: "environment") '
    'untuk memudahkan pemindaian kode cetak. Scanner mendukung format QR Code, Code 128, EAN-13, '
    'dan Code 39. Area pemindaian ditandai dengan kotak bergaris yang berkedip untuk memandu '
    'pengguna.'
)

add_code(
    'useEffect(() => {\n'
    '  html5QrCodeRef.current = new Html5Qrcode(scannerId);\n'
    '  const startCamera = async () => {\n'
    '    await html5QrCodeRef.current?.start(\n'
    '      { facingMode: "environment" },\n'
    '      { fps: 15, qrbox: { width: boxSize, height: boxSize },\n'
    '        formatsToSupport: [\n'
    '          Html5QrcodeSupportedFormats.QR_CODE,\n'
    '          Html5QrcodeSupportedFormats.CODE_128,\n'
    '        ]\n'
    '      },\n'
    '      (decodedText) => { onScan(decodedText); },\n'
    '      () => {}\n'
    '    );\n'
    '  };\n'
    '  startCamera();\n'
    '}, []);'
)
add_figure_caption('[Gambar 4.13 Scanner Kamera]')

doc.add_heading('4.5.3 ScanInterface.tsx', level=3)
add_para(
    'Scan Interface merupakan antarmuka utama untuk petugas Security. Terdapat input teks untuk '
    'pemindaian manual (menggunakan scanner hardware keyboard wedge) dan tombol untuk mengaktifkan '
    'kamera. Sistem menerapkan mekanisme cooldown 10 detik di sisi klien untuk mencegah pemindaian '
    'ganda. Di bagian footer, terdapat grid statistik real-time yang menampilkan data IN, OUT, POB, '
    'dan Visitor yang diperbarui secara otomatis menggunakan onSnapshot Firestore.'
)

add_code(
    'const handleScan = async (nik: string) => {\n'
    '  const cleanNik = nik.replace(/[\\u0000-\\u001F\\u007F-\\u009F]/g, "").trim();\n'
    '  if (!cleanNik || isProcessingRef.current) return;\n'
    '  const now = Date.now();\n'
    '  if (now - (recentScansRef.current[cleanNik] || 0) < 10000) {\n'
    '    setNikInput(""); return;\n'
    '  }\n'
    '  isProcessingRef.current = true;\n'
    '  recentScansRef.current[cleanNik] = now;\n'
    '  setIsProcessing(true);\n'
    '  setNikInput("");\n'
    '  const result = await processScan(cleanNik);\n'
    '  // Handle result...\n'
    '  isProcessingRef.current = false;\n'
    '  setIsProcessing(false);\n'
    '};'
)
add_figure_caption('[Gambar 4.10 Halaman Scan Interface]')

doc.add_heading('4.5.4 AdminDashboard.tsx', level=3)
add_para(
    'Admin Dashboard menyediakan dua tab utama: Attendance Log dan Employee Database. Tab Attendance '
    'Log menampilkan daftar log kehadiran yang dapat difilter berdasarkan tanggal dan departemen. Tab '
    'Employee Database menampilkan daftar personel dengan status kehadiran terkini (IN/OUT) serta '
    'tombol aksi untuk mengelola data (edit, delete, cetak barcode). Fitur unggulan meliputi:'
)
add_bullet('CRUD Karyawan dan Visitor melalui form modal')
add_bullet('Export PDF laporan kehadiran menggunakan jsPDF dan jspdf-autotable')
add_bullet('Seed Data 90 data karyawan dummy untuk pengujian')
add_bullet('Toggle Maintenance Mode untuk membatasi akses non-admin')
add_bullet('Cetak Identity Pass berisi QR Code dan Barcode')

add_code(
    'const handleExport = () => {\n'
    '  const doc = new jsPDF();\n'
    '  doc.setFontSize(22);\n'
    '  doc.text("Laporan Kehadiran Warehouse", 14, 20);\n'
    '  doc.setFontSize(10);\n'
    '  doc.text(`Warehouse ELNUSA BSD - ${selectedDate}`, 14, 28);\n'
    '  const tableData = filteredLogs.map(log => [\n'
    '    employees[log.employeeId]?.name || "Unknown",\n'
    '    log.employeeId,\n'
    '    employees[log.employeeId]?.department || "N/A",\n'
    '    log.timestamp ? format(log.timestamp, "HH:mm:ss") : "N/A",\n'
    '    log.type\n'
    '  ]);\n'
    '  autoTable(doc, {\n'
    '    startY: 48,\n'
    '    head: [["Nama", "NIK", "Departemen", "Waktu", "Status"]],\n'
    '    body: tableData,\n'
    '    theme: "grid"\n'
    '  });\n'
    '  doc.save(`Warehouse_Attendance_${selectedDate}.pdf`);\n'
    '};'
)
add_figure_caption('[Gambar 4.11 Halaman Admin Dashboard]')

doc.add_heading('4.5.5 attendance.ts (Logika Absensi)', level=3)
add_para(
    'File attendance.ts berisi logika inti sistem absensi. Fungsi getEmployeeByNik menggunakan '
    'enam strategi pencarian berurutan: (1) Direct Document ID, (2) Exact NIK Field, (3) Case-'
    'Insensitive Search, (4) Numeric Normalization, (5) Name Search, dan (6) Partial Name Match. '
    'Fungsi processScan mengelola toggle IN/OUT secara otomatis berdasarkan log terakhir, serta '
    'melakukan batch write atomik menggunakan writeBatch Firestore untuk mencatat log dan '
    'memperbarui statistik secara bersamaan.'
)

add_code(
    'export async function processScan(nik: string) {\n'
    '  const employee = await getEmployeeByNik(nik);\n'
    '  if (!employee) return { success: false, message: "Karyawan tidak ditemukan" };\n'
    '  const latestLog = await getLatestLog(employee.id, format(new Date(), "yyyy-MM-dd"));\n'
    '  // Cooldown 1 menit\n'
    '  if (latestLog?.timestamp && (Date.now() - latestLog.timestamp.seconds*1000) < 60000)\n'
    '    return { success: false, message: "Mohon tunggu sebentar", employee };\n'
    '  let nextType = PresenceType.IN;\n'
    '  if (latestLog?.type === PresenceType.IN) nextType = PresenceType.OUT;\n'
    '  const batch = writeBatch(db);\n'
    '  // Write log + update stats atomically\n'
    '  batch.set(doc(collection(db, "presence_logs")), {\n'
    '    employeeId: employee.id, type: nextType,\n'
    '    timestamp: serverTimestamp(), date: todayStr\n'
    '  });\n'
    '  const statsUpdate = nextType === PresenceType.IN\n'
    '    ? { in: increment(1), pob: increment(1) }\n'
    '    : { out: increment(1), pob: increment(-1) };\n'
    '  batch.set(doc(db, "stats", todayStr), statsUpdate, { merge: true });\n'
    '  await batch.commit();\n'
    '  return { success: true, message: msg, employee, type: nextType };\n'
    '}'
)

doc.add_heading('4.5.6 App.tsx (Main App)', level=3)
add_para(
    'App.tsx merupakan komponen utama yang mengelola routing aplikasi berdasarkan peran pengguna. '
    'Jika belum login, ditampilkan LoginSelection. Jika login sebagai Employee, ditampilkan portal '
    'karyawan yang berisi QR Code pribadi, Barcode, dan riwayat kehadiran. Jika login sebagai Admin '
    'atau Security, ditampilkan tombol navigasi antara ScanInterface dan AdminDashboard. App.tsx '
    'juga mengelola fitur Maintenance Mode yang memblokir akses non-admin ketika diaktifkan.'
)
add_figure_caption('[Gambar 4.12 Halaman Employee Portal]')

doc.add_heading('4.5.7 Implementasi PWA (Progressive Web App)', level=3)
add_para(
    'Progressive Web App (PWA) diimplementasikan untuk memenuhi kebutuhan akses mobile tanpa '
    'harus mengembangkan aplikasi native terpisah. PWA memungkinkan aplikasi web untuk diinstal '
    'pada perangkat Android dan iOS sehingga dapat berjalan seperti aplikasi native, termasuk '
    'kemampuan offline, notifikasi push, dan akses layar penuh tanpa browser URL bar.'
)
add_para(
    'Implementasi PWA menggunakan library vite-plugin-pwa yang terintegrasi langsung dengan '
    'Vite build tool. Konfigurasi dilakukan pada file vite.config.ts dengan pengaturan sebagai '
    'berikut:'
)
add_code(
    "VitePWA({\n"
    "  registerType: 'autoUpdate',\n"
    '  includeAssets: [\'favicon.svg\', \'pwa-icon.svg\', \'icon-180.png\', \'icon-120.png\'],\n'
    '  manifest: {\n'
    "    name: 'Warehouse Elnusa BSD - Presence System',\n"
    "    short_name: 'ElnusaAbsensi',\n"
    "    description: 'Warehouse ELNUSA BSD Integrated Presence & Warehouse Management System',\n"
    "    theme_color: '#2563eb',\n"
    "    background_color: '#020617',\n"
    "    display: 'standalone',\n"
    "    orientation: 'any',\n"
    "    scope: '/',\n"
    "    start_url: '/',\n"
    '    icons: [\n'
    "      { src: 'pwa-icon.svg', sizes: 'any', type: 'image/svg+xml' },\n"
    "      { src: 'icon-120.png', sizes: '120x120', type: 'image/png' },\n"
    "      { src: 'icon-180.png', sizes: '180x180', type: 'image/png' },\n"
    '    ],\n'
    '  },\n'
    '  workbox: {\n'
    "    globPatterns: ['**/*.{js,css,html,svg,png,ico,woff2}'],\n"
    '  },\n'
    '})'
)
add_para(
    'Service worker dibangun menggunakan Workbox (library Google untuk service worker) dengan '
    'strategi precache yang menyimpan seluruh aset statis (JS, CSS, HTML, SVG, PNG) sebanyak '
    '16 entri dengan total ~2.1 MB. Hal ini memungkinkan aplikasi tetap dapat diakses meskipun '
    'dalam kondisi jaringan terbatas atau offline.'
)
add_para(
    'Untuk mendukung perangkat iOS (iPhone/iPad), ditambahkan meta tags khusus Safari pada '
    'file index.html:'
)
add_code(
    '<meta name="apple-mobile-web-app-capable" content="yes" />\n'
    '<meta name="apple-mobile-web-app-status-bar-style" content="black-translucent" />\n'
    '<meta name="apple-mobile-web-app-title" content="ElnusaAbsensi" />\n'
    '<link rel="apple-touch-icon" sizes="180x180" href="/icon-180.png" />\n'
    '<link rel="apple-touch-icon" sizes="120x120" href="/icon-120.png" />'
)
add_para(
    'Dengan implementasi PWA ini, aplikasi dapat diakses melalui menu "Add to Home Screen" '
    'pada browser Safari (iOS) dan Chrome (Android), sehingga pengguna (petugas Security dan '
    'Admin) dapat menjalankan sistem absensi langsung dari layar utama perangkat mobile mereka '
    'tanpa perlu menginstal aplikasi dari App Store atau Play Store.'
)
add_image('/root/ElnusaAbsensiWEB/public/icon-180.png')
add_figure_caption('Gambar 4.14 Icon Aplikasi PWA (180x180 pixel)')
add_image('/root/ElnusaAbsensiWEB/public/icon-120.png')
add_figure_caption('Gambar 4.15 Icon Aplikasi PWA Ukuran 120x120 pixel')

doc.add_heading('4.5.8 WarehouseDashboard.tsx (POB Monitoring)', level=3)
add_para(
    'WarehouseDashboard merupakan komponen monitoring yang menampilkan dashboard analitik '
    'Person On Board (POB) secara real-time. Komponen ini menggunakan library Recharts untuk '
    'visualisasi data dalam bentuk Pie Chart (donut) dan Bar Chart (trend). Dashboard terintegrasi '
    'dengan Firestore melalui onSnapshot untuk menampilkan jumlah personel yang sedang aktif di '
    'lokasi (POB) secara langsung.'
)
add_para(
    'Fitur-fitur yang terdapat pada WarehouseDashboard meliputi:'
)
add_bullet('Filter tanggal kehadiran dengan rentang waktu yang dapat dipilih (date range picker).')
add_bullet('Pie Chart donut yang menampilkan komposisi POB berdasarkan kategori personel.')
add_bullet('Bar Chart trend kehadiran harian untuk 14 hari terakhir.')
add_bullet('Tabel breakdown POB per kategori personel dengan total akumulatif.')
add_bullet('Panel breakdown POB per shift kerja (Shift 1, Shift 2, Shift 3, Office).')
add_bullet('Peta mini Warehouse BSD yang menunjukkan lokasi personel aktif.')
add_bullet('Panduan integrasi API untuk mengganti data mock dengan data real-time dari database.')
add_para(
    'Dashboard ini dapat diakses oleh Admin melalui tombol "POB Monitoring Dashboard" pada '
    'navigasi floating di sudut kanan bawah aplikasi. Untuk pengalaman mobile yang optimal, '
    'dashboard diatur dengan tata letak responsif yang menumpuk panel secara vertikal pada '
    'perangkat dengan layar kecil.'
)
add_figure_caption('[Gambar 4.16 Tampilan WarehouseDashboard POB Monitoring]')

doc.add_heading('4.6 User Interface Website', level=2)

doc.add_heading('Halaman Login', level=3)
add_para(
    'Halaman login menampilkan tiga kartu pilihan yang merepresentasikan tiga peran pengguna dengan '
    'desain modern menggunakan gradasi warna dan animasi. Setiap kartu memiliki ikon, deskripsi peran, '
    'dan tombol akses yang sesuai. Modal login muncul dengan animasi smooth saat tombol diklik.'
)
add_figure_caption('[Gambar 4.9 Halaman Login (LoginSelection)]')

doc.add_heading('Halaman Scan Interface', level=3)
add_para(
    'Halaman Scan Interface menampilkan input teks untuk pemindaian NIK manual dan tombol kamera '
    'untuk pemindaian QR/Barcode. Di bagian bawah terdapat grid statistik real-time yang menampilkan '
    'jumlah personel masuk (IN), keluar (OUT), Person On Board (POB), dan data visitor. Notifikasi '
    'hasil scan muncul dalam bentuk animasi yang informatif.'
)
add_figure_caption('[Gambar 4.10 Halaman Scan Interface]')

doc.add_heading('Halaman Admin Dashboard', level=3)
add_para(
    'Halaman Admin Dashboard menampilkan data log kehadiran dan manajemen karyawan dalam dua tab. '
    'Dilengkapi dengan filter tanggal dan departemen, tombol ekspor PDF, tombol seed data, serta '
    'tombol toggle maintenance mode. Sidebar menampilkan status kehadiran seluruh personel secara '
    'real-time.'
)
add_figure_caption('[Gambar 4.11 Halaman Admin Dashboard]')

doc.add_heading('Halaman Employee Portal', level=3)
add_para(
    'Halaman Employee Portal menampilkan QR Code dan Barcode pribadi karyawan, riwayat kehadiran '
    'dalam bentuk tabel yang diurutkan berdasarkan waktu terbaru, serta ringkasan statistik pribadi '
    'seperti total scan dan jumlah kehadiran hari ini.'
)
add_figure_caption('[Gambar 4.12 Halaman Employee Portal]')

doc.add_heading('Halaman PWA Install di Perangkat Mobile', level=3)
add_para(
    'Setelah implementasi PWA, pengguna dapat menginstal aplikasi ke layar utama perangkat mobile '
    'melalui menu "Add to Home Screen" atau "Install App" yang muncul secara otomatis di browser '
    'Chrome (Android) dan Safari (iOS). Saat diinstal, aplikasi berjalan dalam mode standalone '
    '(layar penuh tanpa URL bar browser) dengan icon aplikasi dan splash screen khas aplikasi '
    'native. Proses instalasi PWA di iPhone Safari meliputi: (1) membuka tautan aplikasi di '
    'Safari, (2) menekan tombol Share (ikon kotak dengan panah), (3) memilih "Add to Home Screen", '
    '(4) menekan "Add" pada pojok kanan atas.'
)
add_para(
    'Setelah terinstal, aplikasi hadir sebagai ikon di layar utama iPhone dengan tampilan layar '
    'penuh tanpa browser chrome, mendukung orientasi potrait dan landscape, serta memberikan '
    'pengalaman yang hampir identik dengan aplikasi native. Hal ini membuktikan bahwa konsep '
    'Pemrograman Mobile (CIE407) telah berhasil diterapkan tanpa perlu mengembangkan aplikasi '
    'Android native atau iOS native.'
)
add_figure_caption('[Gambar 4.17 Tampilan PWA Terinstal di Layar Utama iPhone]')

doc.add_heading('4.7 Pemetaan dan Bukti (Evidence) Konversi Mata Kuliah', level=2)

t3 = doc.add_table(rows=1, cols=4)
t3.style = 'Table Grid'
for i, h in enumerate(['No', 'Mata Kuliah', 'Penerapan dalam Kegiatan Magang', 'Bukti (Evidence)']):
    t3.rows[0].cells[i].text = h
    for r in t3.rows[0].cells[i].paragraphs:
        for run in r.runs:
            run.bold = True
            run.font.size = Pt(9)

mk_data = [
    ('1', 'Arsitektur Berbasis Layanan (CIE408)',
     'Mengimplementasikan arsitektur berbasis layanan dengan React component-based architecture pada frontend, '
     'Express REST API pada sisi backend server, serta Firebase Firestore sebagai service layer basis data. '
     'Komunikasi antar-layanan menggunakan protokol HTTP untuk request-response dan real-time listener '
     'menggunakan onSnapshot Firestore untuk pembaruan data secara langsung.',
     'App.tsx (routing komponen), server.ts (Express API), firebase.ts dan attendance.ts (Firestore service layer)'),
    ('2', 'Isu Sosial & Profesional TI (CIE617)',
     'Menerapkan aspek sosial dan profesionalisme TI melalui sistem role-based access control (RBAC) dengan '
     'tiga peran pengguna (Admin, Security, Employee) yang memiliki hak akses berbeda. Mekanisme cooldown '
     '10 detik untuk mencegah penyalahgunaan scan, fitur maintenance mode untuk pengendalian akses saat '
     'pemeliharaan, serta pengelolaan data personel secara etis, akurat, dan transparan.',

     'types.ts (UserRole enum), App.tsx (role-based routing), attendance.ts (cooldown logic), system_config (maintenance mode)'),
    ('3', 'Interaksi Manusia Komputer (CSF619)',
     'Merancang antarmuka yang intuitif dan responsif menggunakan Tailwind CSS dengan prinsip-prinsip '
     'interaksi manusia dan komputer. Fitur interaktif meliputi notifikasi animatif hasil scan, umpan '
     'balik visual real-time (loading spinner, error alert, success toast), animasi transisi halaman, '
     'serta desain mobile-first yang adaptif terhadap berbagai ukuran layar.',
     'ScanInterface.tsx (notifikasi animatif), LoginSelection.tsx (kartu interaktif), index.css (styling responsif)'),
    ('4', 'Pemograman Mobile (CIE407)',
     'Mengimplementasikan Progressive Web App (PWA) menggunakan vite-plugin-pwa yang terintegrasi dengan '
     'Vite build tool. Menghasilkan manifest.json untuk konfigurasi instalasi, service worker dengan '
     'Workbox untuk precache 16 entri aset statis (~2.1 MB), serta meta tags iOS untuk dukungan '
     'Safari pada iPhone/iPad. Aplikasi dapat diinstal melalui "Add to Home Screen" dan berjalan '
     'dalam mode standalone layar penuh tanpa browser URL bar.',
     'vite.config.ts (konfigurasi VitePWA), index.html (iOS meta tags), public/sw.js, public/icon-180.png, public/icon-120.png'),
    ('5', 'Sistem Basis Data Terdistribusi (CIE721)',
     'Merancang dan mengimplementasikan basis data NoSQL terdistribusi menggunakan Firebase Firestore '
     'dengan empat koleksi utama: employees (data personel), presence_logs (catatan kehadiran), stats '
     '(statistik harian), dan system_config (konfigurasi sistem). Memanfaatkan real-time synchronization '
     'melalui onSnapshot untuk pembaruan data langsung, serta batch write atomik menggunakan writeBatch '
     'untuk menjaga konsistensi data pada operasi pencatatan kehadiran.',
     'firebase.ts (inisialisasi Firestore), attendance.ts (query dan batch write), firestore.rules (keamanan data)'),
    ('6', 'Magang (CSF721)',
     'Melaksanakan seluruh rangkaian kegiatan magang di Warehouse Elnusa BSD selama periode April hingga '
     'Agustus 2025 pada posisi IT & ARP. Kegiatan meliputi analisis kebutuhan sistem absensi, perancangan '
     'arsitektur, implementasi kode (React + TypeScript + Firebase), pengujian Black Box, deployment ke '
     'Vercel, serta penyusunan laporan akhir magang ini sebagai dokumentasi lengkap seluruh aktivitas.',
     'Seluruh repositori GitHub (github.com/Adamputra77/ElnusaAbsensiWEB), laporan magang, dan dokumentasi teknis'),
]
for md in mk_data:
    add_table_row(t3, list(md), size=9)

add_figure_caption('Tabel 4.2 Pemetaan dan Bukti (Evidence) Konversi Mata Kuliah')

doc.add_page_break()

# ============================================================
# BAB 5 - PENUTUP
# ============================================================
doc.add_heading('BAB 5\nPENUTUP', level=1)

doc.add_heading('5.1 Kesimpulan', level=2)
add_para(
    'Pelaksanaan Praktik Kerja Lapangan melalui program MBKM Mandiri di Warehouse Elnusa BSD '
    'telah memberikan pengalaman berharga dalam pengembangan sistem informasi di lingkungan industri. '
    'Berdasarkan hasil pelaksanaan magang dan pengembangan sistem, dapat disimpulkan beberapa hal '
    'sebagai berikut:'
)
add_numbered(
    'Sistem absensi berbasis QR Code dan Barcode berhasil dikembangkan menggunakan React.js dengan '
    'TypeScript pada sisi frontend dan Firebase Firestore sebagai basis data real-time. Sistem ini '
    'mampu menggantikan pencatatan kehadiran manual menjadi digital yang lebih efisien dan akurat.'
)
add_numbered(
    'Sistem berhasil mengimplementasikan tiga peran pengguna (Admin, Security, Employee) dengan '
    'hak akses yang berbeda sesuai dengan kebutuhan operasional Warehouse Elnusa BSD.'
)
add_numbered(
    'Integrasi pemindaian QR Code dan Barcode berhasil diimplementasikan melalui dua metode: input '
    'keyboard dari scanner hardware dan pemindaian kamera menggunakan library html5-qrcode dengan '
    'dukungan multi-format.'
)
add_numbered(
    'Fitur-fitur pendukung seperti statistik real-time, ekspor PDF, manajemen data karyawan, '
    'maintenance mode, dan identitas digital (QR Code & Barcode) telah berfungsi dengan baik '
    'berdasarkan pengujian Black Box yang dilakukan.'
)
add_numbered(
    'Sistem berhasil diimplementasikan sebagai Progressive Web App (PWA) menggunakan vite-plugin-pwa '
    'dengan service worker berbasis Workbox yang melakukan precache aset statis, manifest.json untuk '
    'konfigurasi instalasi, serta meta tags iOS untuk dukungan Safari. Aplikasi dapat diinstal melalui '
    '"Add to Home Screen" pada perangkat Android dan iOS, membuktikan penerapan konsep Pemrograman '
    'Mobile tanpa pengembangan aplikasi native.'
)
add_numbered(
    'Warehouse Dashboard monitoring POB (Person On Board) berhasil dikembangkan menggunakan Recharts '
    'dengan visualisasi Pie Chart dan Bar Chart, filter tanggal, serta integrasi real-time Firestore. '
    'Dashboard ini memberikan kemampuan monitoring kehadiran personel secara menyeluruh bagi Admin '
    'melalui tampilan yang responsif dan dapat diakses dari perangkat mobile maupun desktop.'
)

doc.add_heading('5.2 Saran', level=2)
add_para(
    'Berdasarkan pengalaman dan hasil pengembangan yang telah dilakukan, terdapat beberapa saran '
    'yang dapat disampaikan untuk pengembangan sistem selanjutnya:'
)

add_bullet(
    'PWA yang telah diimplementasikan dapat ditingkatkan dengan penambahan fitur push notification '
    'untuk notifikasi real-time tanpa harus membuka aplikasi, serta background sync untuk mendukung '
    'pencatatan kehadiran secara offline yang akan tersinkronisasi secara otomatis saat perangkat '
    'terhubung ke internet kembali.'
)
add_bullet(
    'Penambahan fitur autentikasi berbasis biometrik (wajah atau sidik jari) untuk meningkatkan '
    'keamanan dan mencegah penyalahgunaan identitas.'
)
add_bullet(
    'Integrasi sistem dengan perangkat IoT seperti gate otomatis yang terhubung dengan hasil scan '
    'absensi untuk meningkatkan otomatisasi di pintu gerbang.'
)
add_bullet(
    'Bagi mahasiswa yang akan melaksanakan magang selanjutnya, disarankan untuk mempersiapkan diri '
    'dengan mempelajari teknologi yang akan digunakan di tempat magang, serta aktif berkomunikasi '
    'dengan pembimbing lapangan selama proses magang berlangsung.'
)

doc.add_page_break()

# ============================================================
# DAFTAR PUSTAKA
# ============================================================
doc.add_heading('DAFTAR PUSTAKA', level=1)
add_para('')

references = [
    'Alief, R. N., & Rianto, H. (2025). Perancangan Sistem Absensi Siswa Berbasis Quick Response (QR) Code Menggunakan Framework JavaScript. INSANtek, 6(2), 89-97. https://doi.org/10.31294/insantek.v6i2.10185',
    'Balogun, F. (2026). Design and Implementation of an Enhanced QR-Code Based Attendance System. Journal of Computing and Social Informatics. https://doi.org/10.33736/jcsi.10752.2026',
    'Djamarullah, A. R., Nuryasin, I., & Wibowo, H. (2024). Designing a QR Code Attendance System Using BYOD (Bring Your Own Device). Ultimatics: Jurnal Teknik Informatika, 16(1), 32-37.',
    'Express.js. (2024). Express.js Documentation. https://expressjs.com/',
    'Facebook Open Source. (2024). React Documentation. https://react.dev/',
    'Fadhilah, L. N., Auliana, S., & Aryanto, G. D. P. (2025). Perancangan Sistem Absensi Siswa Berbasis Web Menggunakan QR Code Disekolah PAUD Amelia Darul Akhyar Cikande. Jurnal Multimedia dan Teknologi Informasi (Jatilima), 7(02), 282-290. https://doi.org/10.54209/jatilima.v7i02.1530',
    'Google. (2024). Firebase Documentation. https://firebase.google.com/docs',
    'Hall, H. (2024). jsPDF Library Documentation. https://github.com/parallax/jsPDF',
    'Minh, T. (2024). html5-qrcode Library Documentation. https://github.com/mebjas/html5-qrcode',
    'Mozilla. (2024). MDN Web Docs: QR Code. https://developer.mozilla.org/',
    'Müller, D. (2024). motion/react Animation Library. https://motion.dev/',
    'Nuralif, I., & Fachrie, M. (2023). Development of a QR code-based attendance system for factory employees. International Journal Software Engineering and Computer Science (IJSECS), 3(3), 281-286.',
    'Paz, J. (2024). date-fns Documentation. https://date-fns.org/',
    'Praba, A. D., Safitri, M., & Faridi, F. (2025). Aplikasi Absensi Berbasis Website Menggunakan QR Code untuk Peningkatan Efisiensi Pencatatan Kehadiran. Jurnal Teknik, 14(2). https://doi.org/10.31000/jt.v14i2.15517',
    'Simek, M. (2024). jspdf-autotable Documentation. https://github.com/simonbengtsson/jspdf-autotable',
    'Tailwind Labs. (2024). Tailwind CSS Documentation. https://tailwindcss.com/docs',
    'Cristea, C. (2024). lucide-react Icons Library. https://lucide.dev/',
    'Google Developers. (2024). Workbox: JavaScript Libraries for Progressive Web Apps. https://developer.chrome.com/docs/workbox',
    'Recharts. (2024). Recharts: A composable charting library for React. https://recharts.org/',
    'You, E. (2024). Vite Documentation. https://vitejs.dev/',
]
for i, ref in enumerate(references, 1):
    add_para(f'{i}. {ref}', first_line_indent=False, size=11)

doc.add_page_break()

# ============================================================
# LAMPIRAN (Kosong)
# ============================================================
doc.add_heading('LAMPIRAN 1\nSURAT PERMOHONAN MAGANG', level=1)
add_para('[Halaman ini siap diisi dengan Surat Permohonan Magang]')

doc.add_page_break()
doc.add_heading('LAMPIRAN 2\nSURAT PENERIMAAN MAGANG', level=1)
add_para('[Halaman ini siap diisi dengan Surat Penerimaan Magang]')

doc.add_page_break()
doc.add_heading('LAMPIRAN 3\nLEARNING AGREEMENT', level=1)
add_para('[Halaman ini siap diisi dengan Surat Learning Agreement]')

doc.add_page_break()
doc.add_heading('LAMPIRAN 4\nSERTIFIKAT DAN PENILAIAN MAGANG', level=1)
add_para('[Halaman ini siap diisi dengan Sertifikat dan Penilaian Magang]')

doc.add_page_break()
doc.add_heading('LAMPIRAN 5\nDOKUMENTASI MAGANG', level=1)
add_para('[Halaman ini siap diisi dengan Dokumentasi Magang]')

doc.add_page_break()
doc.add_heading('LAMPIRAN 6\nRIWAYAT HIDUP', level=1)
add_para('[Halaman ini siap diisi dengan Riwayat Hidup]')

# ============================================================
# SAVE
# ============================================================
output_path = '/root/ElnusaAbsensiWEB/LAPORAN AKHIR MAGANG - Adam Putra Pratama.docx'
doc.save(output_path)
print(f"Laporan berhasil disimpan: {output_path}")
print(f"Ukuran file: {os.path.getsize(output_path) / 1024:.1f} KB")
